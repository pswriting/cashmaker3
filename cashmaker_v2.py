import streamlit as st
import google.generativeai as genai
import re
import json
import io
import os
from datetime import datetime
from pathlib import Path

# ==========================================
# API 키 저장/불러오기 (로컬 파일)
# ==========================================
def get_config_path():
    """설정 파일 경로 반환"""
    home = Path.home()
    return home / ".ebook_app_config.json"

def load_saved_api_key():
    """저장된 API 키 불러오기"""
    config_path = get_config_path()
    try:
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
                return config.get('api_key', '')
    except Exception:
        pass
    return ''

def save_api_key(api_key):
    """API 키 저장"""
    config_path = get_config_path()
    try:
        config = {}
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
        config['api_key'] = api_key
        with open(config_path, 'w') as f:
            json.dump(config, f)
        return True
    except Exception:
        return False

# --- 페이지 설정 ---
st.set_page_config(
    page_title="전자책 작성 프로그램", 
    layout="wide", 
    page_icon="◆"
)

# --- 지구인사이트 스타일 CSS ---
st.markdown("""
<style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    
    * { 
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, system-ui, sans-serif; 
    }
    
    /* 기본 요소 숨김 */
    .stDeployButton {display:none;} 
    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    
    /* 사이드바 토글 버튼 표시 */
    [data-testid="collapsedControl"] {
        display: flex !important;
        visibility: visible !important;
    }
    
    /* 메인 배경 - 순수 화이트 */
    .stApp {
        background: #ffffff;
    }
    
    /* 메인 영역 */
    .main .block-container {
        background: #ffffff;
        padding: 2rem 3rem;
        max-width: 1200px;
    }
    
    /* 사이드바 */
    [data-testid="stSidebar"] {
        background: #ffffff;
        border-right: 1px solid #eeeeee;
    }
    
    [data-testid="stSidebar"] * {
        color: #222222 !important;
    }
    
    [data-testid="stSidebar"] .stProgress > div > div > div > div {
        background: #222222;
        border-radius: 10px;
    }
    
    /* 모든 텍스트 - 진한 검정 */
    .stMarkdown, .stText, p, span, label, .stMarkdown p {
        color: #222222 !important;
        line-height: 1.7;
    }
    
    /* 제목 스타일링 */
    h1 { 
        color: #111111 !important; 
        font-weight: 700 !important; 
        font-size: 2rem !important;
        letter-spacing: -0.5px;
        margin-bottom: 1rem !important;
    }
    
    h2 { 
        color: #111111 !important; 
        font-weight: 700 !important;
        font-size: 1.4rem !important;
        margin-top: 2rem !important;
        margin-bottom: 1rem !important;
    }
    
    h3 { 
        color: #222222 !important; 
        font-weight: 600 !important;
        font-size: 1.1rem !important;
        margin-bottom: 0.8rem !important;
    }
    
    /* 탭 스타일 - 미니멀 라인 */
    .stTabs [data-baseweb="tab-list"] {
        background: transparent;
        gap: 0;
        border-bottom: 2px solid #eeeeee;
        padding: 0;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: transparent;
        color: #888888 !important;
        border-radius: 0;
        font-weight: 500;
        padding: 16px 24px;
        font-size: 15px;
        border-bottom: 2px solid transparent;
        margin-bottom: -2px;
        transition: all 0.2s;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        color: #222222 !important;
    }
    
    .stTabs [aria-selected="true"] {
        background: transparent !important;
        color: #111111 !important;
        font-weight: 700 !important;
        border-bottom: 2px solid #111111 !important;
    }
    
    /* 버튼 스타일 - 검정 배경 + 흰색 글씨 */
    .stButton > button { 
        width: 100%; 
        border-radius: 30px; 
        font-weight: 600; 
        background: #111111 !important;
        color: #ffffff !important;
        border: none !important;
        padding: 14px 32px;
        font-size: 15px;
        transition: all 0.2s;
        box-shadow: none;
    }
    
    .stButton > button:hover { 
        background: #333333 !important;
        color: #ffffff !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        transform: translateY(-1px);
    }
    
    .stButton > button:active {
        transform: translateY(0);
    }
    
    /* 버튼 내부 텍스트 강제 흰색 */
    .stButton > button p,
    .stButton > button span,
    .stButton > button div,
    .stButton > button * {
        color: #ffffff !important;
    }
    
    /* 다운로드 버튼 */
    .stDownloadButton > button {
        background: #2d5a27 !important;
        color: #ffffff !important;
        border-radius: 30px;
    }
    
    .stDownloadButton > button:hover {
        background: #3d7a37 !important;
    }
    
    .stDownloadButton > button p,
    .stDownloadButton > button span,
    .stDownloadButton > button * {
        color: #ffffff !important;
    }
    
    /* 입력 필드 */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea {
        background: #ffffff !important;
        border: 1px solid #dddddd !important;
        border-radius: 8px !important;
        color: #222222 !important;
        padding: 14px 16px !important;
        font-size: 15px !important;
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: #111111 !important;
        box-shadow: none !important;
    }
    
    .stTextInput > div > div > input::placeholder,
    .stTextArea > div > div > textarea::placeholder {
        color: #aaaaaa !important;
    }
    
    /* 셀렉트박스 */
    .stSelectbox > div > div {
        background: #ffffff !important;
        border: 1px solid #dddddd !important;
        border-radius: 8px !important;
    }
    
    .stSelectbox > div > div > div {
        color: #222222 !important;
    }
    
    /* 메트릭 */
    [data-testid="stMetricValue"] {
        color: #111111 !important;
        font-size: 2rem !important;
        font-weight: 700 !important;
    }
    
    [data-testid="stMetricLabel"] {
        color: #666666 !important;
    }
    
    /* 알림 메시지 */
    .stSuccess {
        background: #f0f9f0 !important;
        border: 1px solid #c8e6c9 !important;
        border-radius: 8px !important;
    }
    .stSuccess p { color: #2e7d32 !important; }
    
    .stWarning {
        background: #fff8e1 !important;
        border: 1px solid #ffecb3 !important;
        border-radius: 8px !important;
    }
    .stWarning p { color: #f57c00 !important; }
    
    .stError {
        background: #ffebee !important;
        border: 1px solid #ffcdd2 !important;
        border-radius: 8px !important;
    }
    .stError p { color: #c62828 !important; }
    
    .stInfo {
        background: #e3f2fd !important;
        border: 1px solid #bbdefb !important;
        border-radius: 8px !important;
    }
    .stInfo p { color: #1565c0 !important; }
    
    /* 구분선 */
    hr {
        border: none !important;
        border-top: 1px solid #eeeeee !important;
        margin: 2rem 0 !important;
    }
    
    /* 프로그레스 바 */
    .stProgress > div > div > div > div {
        background: #222222;
        border-radius: 10px;
    }
    
    /* ===== 커스텀 컴포넌트 ===== */
    
    /* 로그인 화면 */
    .login-container {
        max-width: 400px;
        margin: 100px auto;
        padding: 40px;
        background: #ffffff;
        border: 1px solid #eeeeee;
        border-radius: 20px;
        text-align: center;
    }
    
    .login-title {
        font-size: 28px;
        font-weight: 700;
        color: #111111;
        margin-bottom: 8px;
    }
    
    .login-subtitle {
        font-size: 15px;
        color: #888888;
        margin-bottom: 30px;
    }
    
    /* 히어로 섹션 */
    .hero-section {
        text-align: center;
        padding: 60px 20px;
        margin-bottom: 40px;
    }
    
    .hero-label {
        font-size: 13px;
        font-weight: 600;
        color: #666666;
        letter-spacing: 3px;
        margin-bottom: 16px;
        text-transform: uppercase;
    }
    
    .hero-title {
        font-size: 42px;
        font-weight: 800;
        color: #111111;
        margin-bottom: 16px;
        letter-spacing: -1px;
        line-height: 1.2;
    }
    
    .hero-subtitle {
        font-size: 18px;
        color: #666666;
        font-weight: 400;
    }
    
    /* 섹션 라벨 */
    .section-label {
        font-size: 12px;
        font-weight: 600;
        color: #888888;
        letter-spacing: 2px;
        margin-bottom: 8px;
        text-transform: uppercase;
    }
    
    /* 점수 카드 */
    .score-card {
        background: #f8f8f8;
        border-radius: 20px;
        padding: 50px 40px;
        text-align: center;
    }
    
    .score-number {
        font-size: 80px;
        font-weight: 800;
        color: #111111;
        line-height: 1;
        margin-bottom: 8px;
    }
    
    .score-label {
        color: #888888;
        font-size: 14px;
        font-weight: 500;
    }
    
    /* 상태 배지 */
    .status-badge {
        display: inline-block;
        padding: 8px 20px;
        border-radius: 20px;
        font-weight: 600;
        font-size: 13px;
        margin-top: 20px;
    }
    
    .status-excellent {
        background: #111111;
        color: #ffffff;
    }
    
    .status-good {
        background: #f0f0f0;
        color: #333333;
    }
    
    .status-warning {
        background: #fff3e0;
        color: #e65100;
    }
    
    /* 정보 카드 */
    .info-card {
        background: #f8f8f8;
        border-radius: 16px;
        padding: 24px;
        margin: 16px 0;
    }
    
    .info-card-title {
        font-size: 12px;
        font-weight: 700;
        color: #888888;
        letter-spacing: 1px;
        margin-bottom: 12px;
        text-transform: uppercase;
    }
    
    .info-card p {
        color: #333333 !important;
        font-size: 15px;
        line-height: 1.8;
        margin: 8px 0;
    }
    
    /* 제목 카드 */
    .title-card {
        background: #ffffff;
        border: 1px solid #eeeeee;
        border-radius: 16px;
        padding: 24px;
        margin: 12px 0;
        transition: all 0.2s;
    }
    
    .title-card:hover {
        border-color: #cccccc;
        box-shadow: 0 4px 20px rgba(0,0,0,0.06);
    }
    
    .title-card .card-number {
        font-size: 12px;
        font-weight: 600;
        color: #aaaaaa;
        margin-bottom: 8px;
    }
    
    .title-card .main-title {
        color: #111111;
        font-size: 18px;
        font-weight: 700;
        margin-bottom: 6px;
    }
    
    .title-card .sub-title {
        color: #666666;
        font-size: 14px;
        margin-bottom: 16px;
    }
    
    .title-card .reason {
        color: #444444;
        font-size: 14px;
        padding: 14px 16px;
        background: #f8f8f8;
        border-radius: 10px;
        line-height: 1.6;
    }
    
    /* 점수 아이템 */
    .score-item {
        background: #ffffff;
        border: 1px solid #eeeeee;
        border-radius: 12px;
        padding: 16px 20px;
        margin: 10px 0;
        display: flex;
        justify-content: space-between;
        align-items: center;
    }
    
    .score-item-label {
        color: #333333;
        font-weight: 500;
        font-size: 15px;
    }
    
    .score-item-value {
        color: #111111;
        font-weight: 700;
        font-size: 20px;
    }
    
    .score-item-reason {
        color: #666666;
        font-size: 14px;
        margin-top: 4px;
        line-height: 1.5;
    }
    
    /* 요약 박스 */
    .summary-box {
        background: #f8f8f8;
        border-radius: 12px;
        padding: 20px;
        margin-top: 20px;
    }
    
    .summary-box p {
        color: #333333 !important;
        font-size: 15px;
        line-height: 1.7;
    }
    
    /* 푸터 */
    .premium-footer {
        text-align: center;
        padding: 40px 20px;
        margin-top: 60px;
        border-top: 1px solid #eeeeee;
    }
    
    .premium-footer-text {
        color: #888888;
        font-size: 14px;
    }
    
    .premium-footer-author {
        color: #222222;
        font-weight: 600;
    }
    
    /* 빈 상태 */
    .empty-state {
        text-align: center;
        padding: 60px 20px;
        background: #f8f8f8;
        border-radius: 16px;
    }
    
    .empty-state p {
        color: #888888 !important;
    }
    
    /* 퀵 액션 박스 */
    .quick-action-box {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        border: 1px dashed #dee2e6;
        border-radius: 16px;
        padding: 24px;
        margin: 16px 0;
        text-align: center;
    }
    
    .quick-action-box p {
        color: #495057 !important;
        font-size: 14px;
        margin-bottom: 12px;
    }
    
    /* 모드 선택 라디오 버튼 스타일 */
    .stRadio > div {
        display: flex;
        gap: 16px;
    }
    
    .stRadio > div > label {
        background: #f8f8f8;
        padding: 12px 20px;
        border-radius: 8px;
        cursor: pointer;
        transition: all 0.2s;
    }
    
    .stRadio > div > label:hover {
        background: #eeeeee;
    }
    
    /* 소제목 카드 */
    .subtopic-card {
        background: #ffffff;
        border: 1px solid #e0e0e0;
        border-radius: 12px;
        padding: 16px;
        margin: 8px 0;
    }
    
    /* 추가/삭제 버튼 작게 */
    .small-btn {
        font-size: 12px !important;
        padding: 6px 12px !important;
    }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 비밀번호 설정 (여기서 변경 가능)
# ==========================================
CORRECT_PASSWORD = "cashmaker2024"
# ==========================================

# --- 비밀번호 확인 ---
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if not st.session_state['authenticated']:
    st.markdown("""
    <div class="login-container">
        <div class="login-title">CASHMAKER</div>
        <div class="login-subtitle">전자책 작성 프로그램</div>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        password_input = st.text_input("비밀번호를 입력하세요", type="password", placeholder="비밀번호")
        
        if st.button("입장하기"):
            if password_input == CORRECT_PASSWORD:
                st.session_state['authenticated'] = True
                st.rerun()
            else:
                st.error("비밀번호가 틀렸습니다")
    
    st.stop()

# --- 세션 초기화 ---
default_states = {
    'topic': '',
    'target_persona': '',
    'pain_points': '',
    'one_line_concept': '',
    'outline': [],
    'chapters': {},
    'current_step': 1,
    'market_analysis': '',
    'book_title': '',
    'subtitle': '',
    'topic_score': None,
    'topic_verdict': None,
    'score_details': None,
    'generated_titles': None,
    'outline_mode': 'ai',  # 'ai' 또는 'manual'
}

for key, value in default_states.items():
    if key not in st.session_state:
        st.session_state[key] = value

# --- Gemini 모델은 사이드바에서 설정됨 ---

# --- 사이드바 ---
with st.sidebar:
    st.markdown("### Progress")
    
    progress_items = [
        bool(st.session_state['topic']),
        bool(st.session_state['target_persona']),
        bool(st.session_state['outline']),
        len(st.session_state['chapters']) > 0,
        any(ch.get('content') for ch in st.session_state['chapters'].values()) if st.session_state['chapters'] else False
    ]
    progress = sum(progress_items) / len(progress_items) * 100
    
    st.progress(progress / 100)
    st.caption(f"{progress:.0f}% 완료")
    
    st.markdown("---")
    st.markdown("### Info")
    if st.session_state['topic']:
        st.caption(f"주제: {st.session_state['topic']}")
    if st.session_state['book_title']:
        st.caption(f"제목: {st.session_state['book_title']}")
    if st.session_state['outline']:
        st.caption(f"목차: {len(st.session_state['outline'])}개")
    
    completed_chapters = sum(1 for ch in st.session_state['chapters'].values() if ch.get('content'))
    if completed_chapters:
        st.caption(f"완성: {completed_chapters}개")
    
    st.markdown("---")
    st.markdown("### 💾 저장/불러오기")
    
    # 저장 데이터 구성
    save_data = {
        'topic': st.session_state.get('topic', ''),
        'target_persona': st.session_state.get('target_persona', ''),
        'pain_points': st.session_state.get('pain_points', ''),
        'one_line_concept': st.session_state.get('one_line_concept', ''),
        'outline': st.session_state.get('outline', []),
        'chapters': st.session_state.get('chapters', {}),
        'book_title': st.session_state.get('book_title', ''),
        'subtitle': st.session_state.get('subtitle', ''),
        'market_analysis': st.session_state.get('market_analysis', ''),
        'topic_score': st.session_state.get('topic_score'),
        'topic_verdict': st.session_state.get('topic_verdict'),
        'score_details': st.session_state.get('score_details'),
        'generated_titles': st.session_state.get('generated_titles'),
    }
    
    # JSON 다운로드 버튼
    save_json = json.dumps(save_data, ensure_ascii=False, indent=2)
    file_name = st.session_state.get('book_title', '전자책') or '전자책'
    file_name = re.sub(r'[^\w\s가-힣-]', '', file_name)[:20]
    
    st.download_button(
        "📥 작업 저장하기",
        save_json,
        file_name=f"{file_name}_{datetime.now().strftime('%m%d_%H%M')}.json",
        mime="application/json",
        use_container_width=True
    )
    
    # 불러오기
    uploaded_file = st.file_uploader(
        "📤 작업 불러오기",
        type=['json'],
        label_visibility="collapsed"
    )
    
    if uploaded_file is not None:
        try:
            loaded_data = json.loads(uploaded_file.read().decode('utf-8'))
            
            if st.button("불러오기 적용", use_container_width=True):
                # 세션에 데이터 적용
                for key in ['topic', 'target_persona', 'pain_points', 'one_line_concept', 
                           'outline', 'chapters', 'book_title', 'subtitle', 'market_analysis',
                           'topic_score', 'topic_verdict', 'score_details', 'generated_titles']:
                    if key in loaded_data:
                        st.session_state[key] = loaded_data[key]
                
                st.success("불러오기 완료!")
                st.rerun()
        except Exception as e:
            st.error(f"파일 오류: {e}")
    
    st.markdown("---")
    st.markdown("### API 설정")
    
    # API 키 세션 스테이트 초기화 (저장된 키 불러오기)
    if 'api_key' not in st.session_state:
        saved_key = load_saved_api_key()
        st.session_state['api_key'] = saved_key
    
    # API 키 입력
    api_key_input = st.text_input(
        "Gemini API 키",
        value=st.session_state['api_key'],
        type="password",
        placeholder="AIza...",
        help="Google AI Studio에서 발급받은 API 키를 입력하세요"
    )
    
    # 입력값 세션에 저장 + 파일에도 저장
    if api_key_input and api_key_input != st.session_state['api_key']:
        st.session_state['api_key'] = api_key_input
        if save_api_key(api_key_input):
            st.toast("✅ API 키가 저장되었습니다!", icon="💾")
    elif api_key_input:
        st.session_state['api_key'] = api_key_input
    
    # API 키 발급 안내
    with st.expander("API 키 발급 방법 (무료)"):
        st.markdown("""
        **2분이면 끝!**
        
        1. [Google AI Studio](https://aistudio.google.com/apikey) 접속
        2. Google 계정으로 로그인
        3. **"API 키 만들기"** 클릭
        4. 생성된 키 복사
        5. 위 입력창에 붙여넣기
        
        ✅ 완전 무료  
        ✅ 신용카드 불필요  
        ✅ 분당 15회 요청 가능
        """)
    
    # API 연결 상태 (간소화)
    if not st.session_state.get('api_key'):
        st.caption("⚠️ API 키를 입력하세요")
    else:
        col_status, col_del = st.columns([3, 1])
        with col_status:
            st.caption("✅ API 키 입력됨 (자동 저장)")
        with col_del:
            if st.button("🗑️", key="del_api_key", help="API 키 삭제"):
                st.session_state['api_key'] = ''
                save_api_key('')
                st.rerun()

# --- AI 함수 ---
def get_api_key():
    return st.session_state.get('api_key', '')

def get_auto_save_data():
    """자동 저장용 데이터 생성"""
    return {
        'topic': st.session_state.get('topic', ''),
        'target_persona': st.session_state.get('target_persona', ''),
        'pain_points': st.session_state.get('pain_points', ''),
        'one_line_concept': st.session_state.get('one_line_concept', ''),
        'outline': st.session_state.get('outline', []),
        'chapters': st.session_state.get('chapters', {}),
        'book_title': st.session_state.get('book_title', ''),
        'subtitle': st.session_state.get('subtitle', ''),
        'market_analysis': st.session_state.get('market_analysis', ''),
        'topic_score': st.session_state.get('topic_score'),
        'topic_verdict': st.session_state.get('topic_verdict'),
        'score_details': st.session_state.get('score_details'),
        'generated_titles': st.session_state.get('generated_titles'),
        'saved_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }

def sync_full_outline():
    """현재 session_state의 outline과 chapters를 기반으로 full_outline 재생성"""
    if not st.session_state.get('outline'):
        return
    
    new_full_outline = ""
    for ch in st.session_state['outline']:
        new_full_outline += f"## {ch}\n"
        if ch in st.session_state.get('chapters', {}):
            for st_name in st.session_state['chapters'][ch].get('subtopics', []):
                new_full_outline += f"- {st_name}\n"
        new_full_outline += "\n"
    
    st.session_state['full_outline'] = new_full_outline.strip()
    # 위젯 키는 직접 설정하지 않음 (rerun 시 자동 반영)

def trigger_auto_save():
    """자동 저장 트리거 - 세션에 플래그 설정 + 목차 동기화"""
    sync_full_outline()  # 목차 변경 시 full_outline 동기화
    st.session_state['auto_save_trigger'] = True

def ask_ai(system_role, prompt, temperature=0.7):
    api_key = get_api_key()
    if not api_key:
        return "⚠️ API 키를 먼저 입력해주세요."
    
    try:
        genai.configure(api_key=api_key)
        ai_model = genai.GenerativeModel('models/gemini-2.0-flash')
        generation_config = genai.types.GenerationConfig(temperature=temperature)
        full_prompt = f"""당신은 {system_role}입니다.

{prompt}

한국어로 답변해주세요."""
        response = ai_model.generate_content(full_prompt, generation_config=generation_config)
        return response.text
    except Exception as e:
        return f"오류 발생: {str(e)}"

def analyze_topic_score(topic):
    prompt = f"""'{topic}' 주제의 전자책 적합도를 분석해주세요.

다음 5가지 항목을 각각 0~100점으로 채점하고, 종합 점수와 판정을 내려주세요.

채점 항목:
1. 시장성 (수요가 있는가?)
2. 수익성 (돈을 지불할 의향이 있는 주제인가?)
3. 차별화 가능성 (경쟁에서 이길 수 있는가?)
4. 작성 난이도 (전자책으로 만들기 쉬운가?)
5. 지속성 (오래 팔릴 수 있는가?)

반드시 아래 JSON 형식으로만 답변하세요. 다른 텍스트 없이 JSON만:
{{
    "market": {{"score": 85, "reason": "이유"}},
    "profit": {{"score": 80, "reason": "이유"}},
    "differentiation": {{"score": 75, "reason": "이유"}},
    "difficulty": {{"score": 90, "reason": "이유"}},
    "sustainability": {{"score": 70, "reason": "이유"}},
    "total_score": 80,
    "verdict": "적합" 또는 "보통" 또는 "부적합",
    "summary": "종합 의견 2~3문장"
}}"""
    return ask_ai("전자책 시장 분석가", prompt, temperature=0.3)

def generate_titles_advanced(topic, persona, pain_points):
    prompt = f"""당신은 자청(역행자), 엠제이 드마코(부의 추월차선), 김승호(돈의 속성)급 베스트셀러 작가입니다.
당신이 쓴 책들은 수십만 부가 팔렸고, 제목만으로 서점에서 손이 가게 만드는 마법을 부립니다.

[분석 대상]
주제: {topic}
타겟: {persona}  
타겟의 속마음: {pain_points}

[베스트셀러 제목의 핵심 원칙]

1. "읽는 순간 뒤통수를 맞은 느낌" - 기존 상식을 정면으로 뒤집어라
   - "역행자" → 남들과 반대로 가야 성공한다는 역설
   - "부의 추월차선" → 느린 차선(직장)에서 빠른 차선으로 갈아타라
   
2. "이건 나만 몰랐던 거 아냐?" - 소외감과 긴급함을 동시에 자극
   - 읽지 않으면 뒤처질 것 같은 불안감
   - 남들은 이미 알고 있다는 느낌
   
3. "구체적 숫자는 신뢰를 만든다" - 모호함 제거
   - "나는 4시간만 일한다" - 구체적이라 믿음이 감
   - "31개월 만에 10억" - 실제 숫자가 주는 힘
   
4. "짧을수록 강하다" - 7자 이내 메인 타이틀
   - "역행자" (3자), "돈의 속성" (4자), "부의 추월차선" (5자)

[실제 베스트셀러 제목 레퍼런스]
- "역행자" - 한 단어로 정체성 규정 (자청)
- "부의 추월차선" - 메타포로 욕망 자극
- "돈, 뜨겁게 사랑하고 차갑게 다루어라" - 대비와 긴장감
- "나는 4시간만 일한다" - 상식 파괴 + 구체적 숫자
- "언스크립티드" - 낯선 단어로 호기심 유발
- "망할 용기" - 역설적 표현으로 충격

[절대 금지 - 이런 제목은 절대 쓰지 마세요]
- "비법", "노하우", "성공", "시작하세요", "방법", "전략", "가이드"
- "~하는 법", "~하기", "완벽한", "쉬운", "단계별"
- 물음표로 끝나는 평범한 질문형
- "데이터 기반", "체계적", "효율적" 같은 교과서 표현
- "입문", "기초", "초보자를 위한"

[미션]
위 원칙으로 {topic} 주제의 전자책 제목 5개를 만들어주세요.
평범하면 실패입니다. 서점에서 이 제목을 본 사람이 "뭐지?" 하고 멈춰서서 집어들게 만드세요.
자청의 "역행자"처럼 단 한 단어로 사람의 정체성을 흔들 수 있다면 최고입니다.

형식 (JSON만 출력):
{{
    "titles": [
        {{
            "title": "7자 이내 임팩트 제목",
            "subtitle": "15자 이내 보조 설명",
            "concept": "이 제목의 핵심 컨셉",
            "why_works": "왜 사람들이 이 제목에 끌리는지 심리학적 이유"
        }}
    ]
}}"""
    return ask_ai("베스트셀러 작가", prompt, temperature=0.9)

def generate_concept(topic, persona, pain_points):
    prompt = f"""주제: {topic}
타겟: {persona}
타겟의 고민: {pain_points}

위 주제에 대해 "이 책 안 읽으면 손해"라는 느낌을 주는 한 줄 컨셉 5개를 만들어주세요.

좋은 컨셉의 조건:
- 상식을 정면으로 부정 ("~한다고? 틀렸다")
- 호기심 자극 ("진짜 이유는 따로 있다")
- 구체적 숫자 포함 ("3개월 만에", "상위 1%")

출력 형식 (이 형식만 출력하세요):

1. [한 줄 컨셉]
   → 왜 끌리는가

2. [한 줄 컨셉]
   → 왜 끌리는가

3. [한 줄 컨셉]
   → 왜 끌리는가

4. [한 줄 컨셉]
   → 왜 끌리는가

5. [한 줄 컨셉]
   → 왜 끌리는가"""
    return ask_ai("카피라이터", prompt, temperature=0.9)

def generate_outline(topic, persona, pain_points):
    prompt = f"""주제: {topic}
타겟: {persona}
타겟의 고민: {pain_points}

당신은 크몽에서 수천 권을 판매한 베스트셀러 전자책 기획자입니다.
목차만 보고도 "이거 안 사면 진짜 손해다"라고 느끼게 만드는 목차를 설계해주세요.

[미션]
정확히 4개 챕터, 각 챕터당 3개 소제목 (총 12개)

[★★★ 프드프 킬러 목차의 비밀 ★★★]

1. 뒤통수 치는 반전
   - "열심히 하면 성공한다? 거짓말입니다"
   - "저축하면 부자된다는 착각"
   - "노력이 배신하는 진짜 이유"

2. 숫자로 증명하는 구체성
   - "31개월 만에 1억 만든 공식"
   - "하루 47분으로 월 300 추가"
   - "단 3개만 바꿨더니 수입 2배"

3. 빈칸으로 미치게 만드는 호기심
   - "부자들이 절대 안 알려주는 ○○○"
   - "'이것' 모르면 평생 월급쟁이"
   - "99%가 모르는 '그 방법'"

4. 공포와 긴급함
   - "지금 안 하면 5년 후 후회합니다"
   - "이대로면 당신은 절대 부자 못 됩니다"
   - "남들은 이미 시작했습니다"

5. 비밀 공개 + 내부자 정보
   - "아무도 안 알려주는 진짜 비결"
   - "상위 1%만 아는 숨겨진 법칙"
   - "업계에서 쉬쉬하는 그 방법"

[실제 베스트셀러 목차 레퍼런스]

챕터 제목급:
- "나는 왜 맨날 돈이 없을까"
- "월급쟁이는 절대 부자 될 수 없는 이유"
- "상위 1%가 절대 공개 안 하는 것"
- "이것만 알면 인생이 바뀝니다"

소제목급:
- "통장 잔고 47만원, 그날의 깨달음"
- "회사 그만두고 3개월 후 생긴 일"
- "아내가 먼저 알아챈 변화"
- "'그 습관' 하나로 월 200 추가"
- "직장 동료는 아직도 모릅니다"
- "부모님께 처음으로 용돈 드린 날"

[★★★ 절대 금지 - 이런 제목은 0점 ★★★]
- "~의 중요성", "~하는 방법", "~의 기초"
- "효과적인 ~", "성공적인 ~", "올바른 ~"
- "~란 무엇인가", "~의 이해"
- "1장, 2장" 같은 밋밋한 번호
- "시작하기", "마무리하기", "정리하기"
- 물음표로 끝나는 평범한 질문
- 교과서 같은 딱딱한 표현

[감정선 설계]
챕터1: 💥 충격 + 공감 → "맞아, 이게 내 얘기야"
챕터2: 😤 분노 + 깨달음 → "그래서 안 됐구나!"
챕터3: 🔥 희망 + 비밀 → "이게 답이었어!"
챕터4: 💪 확신 + 행동 → "나도 할 수 있어!"

출력 형식 (정확히 이 형식만):

## 챕터1: [뒤통수 치는 충격적 제목]
- [궁금해서 미칠 것 같은 소제목]
- [안 읽으면 손해인 것 같은 소제목]
- [클릭 안 할 수 없는 소제목]

## 챕터2: [분노 유발 + 깨달음 제목]
- [호기심 폭발 소제목]
- [숫자로 증명하는 소제목]
- [빈칸 호기심 소제목]

## 챕터3: [비밀 공개형 제목]
- [내부자 정보 느낌 소제목]
- [구체적 결과 약속 소제목]
- [반전 있는 소제목]

## 챕터4: [행동 촉구 + 비전 제목]
- [따라하면 되는 느낌 소제목]
- [미래 그림 그리는 소제목]
- [확신 주는 소제목]"""
    return ask_ai("베스트셀러 기획자", prompt, temperature=0.95)

def regenerate_chapter_outline(chapter_number, topic, persona, existing_chapters):
    """특정 챕터만 재생성 - 킬러 목차 스타일"""
    
    chapter_emotions = {
        1: "💥 충격 + 공감: '맞아, 이게 내 얘기야' 느낌",
        2: "😤 분노 + 깨달음: '그래서 안 됐구나!' 느낌",
        3: "🔥 희망 + 비밀: '이게 답이었어!' 느낌",
        4: "💪 확신 + 행동: '나도 할 수 있어!' 느낌"
    }
    emotion = chapter_emotions.get(chapter_number, "호기심 폭발")
    
    prompt = f"""주제: {topic}

{chapter_number}번째 챕터를 새로 만들어주세요.
이 챕터의 감정선: {emotion}

[★★★ 킬러 목차 공식 ★★★]

챕터 제목 (20자 이내) - 이 중 하나 이상 필수:
✓ 뒤통수 치는 반전: "~라고? 거짓말입니다"
✓ 숫자 증명: "31개월 만에 1억", "하루 47분"
✓ 빈칸 호기심: "○○○ 하나로 인생역전"
✓ 공포/긴급: "지금 안 하면 평생 후회"
✓ 비밀 공개: "아무도 안 알려주는 진짜"

소제목 (15자 이내) - 각각 다른 기법 사용:
✓ "통장 잔고 47만원, 그날의 깨달음"
✓ "'그 습관' 하나로 월 200 추가"  
✓ "직장 동료는 아직도 모릅니다"
✓ "3개월 후 아내가 먼저 알아챘습니다"
✓ "99%가 모르는 '그 방법'"

[절대 금지]
- "~의 중요성", "~하는 방법"
- "효과적인", "성공적인", "올바른"
- 교과서 같은 딱딱한 표현
- 물음표로 끝나는 평범한 질문

출력 (정확히 이 형식만):
## 챕터{chapter_number}: [뒤통수 치는 제목]
- [클릭 안 할 수 없는 소제목]
- [숫자로 증명하는 소제목]
- [빈칸 호기심 소제목]"""
    return ask_ai("베스트셀러 기획자", prompt, temperature=0.95)

def regenerate_single_subtopic(chapter_title, subtopic_index, topic, existing_subtopics):
    """특정 소제목만 재생성 - 클릭 유발 스타일"""
    prompt = f"""주제: {topic}
챕터: {chapter_title}

{subtopic_index}번 소제목을 새로 만들어주세요.
목표: 이 소제목만 보고도 "읽어야겠다" 느낌이 들게

[★★★ 클릭 유발 소제목 공식 ★★★] (15자 이내)

1. 스토리 암시형
   - "통장 잔고 47만원, 그날의 깨달음"
   - "퇴사 3개월 후 생긴 일"
   - "아내가 먼저 알아챈 변화"

2. 숫자 증명형
   - "하루 47분으로 월 300 추가"
   - "단 3개만 바꿨더니 수입 2배"
   - "21일 만에 첫 100만원"

3. 빈칸 호기심형
   - "○○○ 하나로 인생역전"
   - "'그것' 모르면 평생 헛수고"
   - "99%가 모르는 '그 방법'"

4. 비밀 공개형
   - "아무도 안 알려주는 진짜 비결"
   - "상위 1%만 아는 습관"
   - "업계에서 쉬쉬하는 것"

5. 비교/대조형
   - "월급 300 vs 부업 300"
   - "실패한 나 vs 성공한 나"
   - "1년 전 vs 지금"

[절대 금지]
- "~의 중요성", "~하는 방법", "~의 이해"
- "효과적인", "성공적인", "올바른"
- 15자 초과

출력 (소제목만, 기호 없이):"""
    result = ask_ai("베스트셀러 기획자", prompt, temperature=0.95)
    # 결과 정제
    result = result.strip().strip('[]').strip('-').strip('"').strip("'").strip()
    # 줄바꿈 있으면 첫 줄만
    if '\n' in result:
        result = result.split('\n')[0].strip()
    # 앞뒤 특수문자 제거
    result = result.lstrip('0123456789.-) ').strip()
    return result

def generate_subtopics(chapter_title, topic, persona, num_subtopics=3):
    prompt = f"""주제: {topic}
챕터: {chapter_title}
타겟: {persona}

이 챕터의 소제목 {num_subtopics}개를 만들어주세요.
목표: 각 소제목만 보고도 "이거 꼭 읽어야겠다" 느낌

[★★★ 클릭률 폭발 소제목 공식 ★★★]

1. 스토리 암시형 (구체적 상황)
   - "통장 잔고 47만원, 그날의 깨달음"
   - "퇴사 3개월 후 생긴 일"
   - "아내가 먼저 알아챈 변화"
   - "그날 새벽 3시, 모든 게 바뀌었다"

2. 숫자 증명형 (구체적 결과)
   - "하루 47분으로 월 300 추가"
   - "단 3개만 바꿨더니 수입 2배"
   - "21일 만에 첫 100만원"
   - "7일간의 실험, 놀라운 결과"

3. 빈칸 호기심형 (궁금증 유발)
   - "○○○ 하나로 인생역전"
   - "'그것' 모르면 평생 헛수고"
   - "99%가 모르는 '그 방법'"
   - "부자들만 아는 ○○의 비밀"

4. 비밀 공개형 (내부자 정보)
   - "아무도 안 알려주는 진짜 비결"
   - "상위 1%만 아는 습관"
   - "업계에서 쉬쉬하는 것"
   - "전문가들이 숨기는 진실"

5. 도발/반전형 (상식 파괴)
   - "열심히 하면 망합니다"
   - "노력이 배신하는 진짜 이유"
   - "저축하면 가난해지는 원리"
   - "성실함이 독이 되는 순간"

6. 비교 대조형 (명확한 차이)
   - "월급 300 vs 부업 300"
   - "실패하는 사람 vs 성공하는 사람"
   - "1년 전 나 vs 지금 나"

[★★★ 절대 금지 ★★★]
- "~의 중요성", "~하는 방법", "~의 이해"
- "효과적인 ~", "성공적인 ~", "올바른 ~"
- "~란 무엇인가", "~시작하기"
- 15자 초과하는 긴 제목
- 교과서처럼 딱딱한 표현
- 물음표로 끝나는 평범한 질문

[중요] 각 소제목은 서로 다른 공식을 사용하세요!

출력 형식 (숫자와 소제목만):
1. 소제목
2. 소제목
3. 소제목"""
    return ask_ai("베스트셀러 기획자", prompt, temperature=0.95)

def generate_interview_questions(subtopic_title, chapter_title, topic):
    prompt = f"""당신은 베스트셀러 작가의 고스트라이터입니다.
'{topic}' 전자책의 '{chapter_title}' 챕터 중 '{subtopic_title}' 소제목 부분을 쓰기 위해 작가를 인터뷰합니다.

[인터뷰 목적]
'{subtopic_title}'에 대한 작가의 진짜 경험과 통찰을 끌어내서, 독자가 "와, 이건 진짜 경험한 사람만 알 수 있는 거다"라고 느끼게 만들 콘텐츠를 확보하는 것.

[좋은 질문의 특징]
1. 구체적 상황을 묻는다: "언제, 어디서, 어떻게"
2. 감정을 묻는다: "그때 기분이 어땠나요?"
3. 실패를 묻는다: "처음에 뭘 잘못했나요?"
4. 반전을 묻는다: "뭘 깨닫고 달라졌나요?"
5. 디테일을 묻는다: "구체적으로 어떻게 했나요?"

[나쁜 질문 예시 - 이런 질문은 피하세요]
- "이것의 중요성에 대해 말씀해주세요" (추상적)
- "팁이 있다면?" (뻔한 답변 유도)
- "어떻게 생각하세요?" (의견만 나옴)

[좋은 질문 예시]
- "처음 이걸 시작했을 때 가장 크게 실패한 경험은 뭔가요? 그때 뭘 잘못 생각했던 건가요?"
- "이걸 깨닫기 전과 후, 구체적으로 뭐가 달라졌나요? 숫자로 말해주실 수 있나요?"
- "주변에서 반대했을 때 어떻게 대응했나요? 실제로 뭐라고 말했나요?"
- "이 방법을 처음 시도한 날, 그 상황을 자세히 묘사해주실 수 있나요?"
- "독자들이 가장 많이 하는 실수는 뭔가요? 왜 그 실수를 하게 되나요?"

[미션]
'{subtopic_title}' 소제목의 핵심 내용을 끌어낼 수 있는 인터뷰 질문 3개를 만들어주세요.
이 질문에 답하면 자연스럽게 이 소제목에 대한 몰입감 있는 내용이 완성될 수 있어야 합니다.

형식:
Q1: [구체적이고 깊이 있는 질문]
Q2: [구체적이고 깊이 있는 질문]
Q3: [구체적이고 깊이 있는 질문]"""
    return ask_ai("베스트셀러 고스트라이터", prompt, temperature=0.7)

def generate_subtopic_content(subtopic_title, chapter_title, questions, answers, topic, persona):
    qa_pairs = ""
    for i, (q, a) in enumerate(zip(questions, answers), 1):
        if a.strip():
            qa_pairs += f"\n질문{i}: {q}\n답변{i}: {a}\n"
    
    prompt = f"""당신은 자청(역행자), 신사임당 수준의 베스트셀러 작가입니다.

[집필 정보]
주제: {topic}
챕터: {chapter_title}
현재 작성할 소제목: {subtopic_title}
타겟: {persona}

⚠️ 매우 중요: 오직 '{subtopic_title}'에 대한 본문만 작성하세요.
- 다른 챕터나 소제목 내용을 언급하지 마세요
- 이 소제목의 핵심 주제에만 집중하세요
- 소제목 제목을 본문에 다시 쓰지 마세요

[작가 인터뷰 - 이 내용만 바탕으로 작성]
{qa_pairs}

[★★★ 문체: 존댓말(합쇼체) 100% 통일 ★★★]

모든 문장 끝을 다음 중 하나로만 마무리하세요:
- ~입니다 / ~습니다 / ~했습니다 / ~됩니다
- ~죠 / ~거죠 / ~셨죠
- ~세요 / ~하세요

절대 금지 (반말):
- ~다 / ~했다 / ~이다 / ~였다 / ~된다
- ~라 / ~했다 / ~인 것이다

[자청 + PS글쓰기 스타일]

1. 첫 문장 = 후킹
   좋은 예시:
   - "월급 230만 원. 그게 제 전부였습니다."
   - "솔직히 말씀드리면, 저도 처음엔 실패했습니다."
   - "3개월 동안 매출 0원이었습니다."
   
   나쁜 예시 (절대 쓰지 마세요):
   - "오늘은 ~에 대해 이야기해보겠습니다."
   - "~의 중요성에 대해 말씀드리겠습니다."

2. 문단 구성 (가독성 핵심)
   - 한 문단 = 3~5문장 묶기
   - 문단과 문단 사이에만 빈 줄 1개
   - 한 문장씩 띄어쓰기 절대 금지

3. 자연스러운 스토리 전개
   - 에피소드 → 깨달음 → 변화 → 결과
   - 구체적 상황 묘사 (날짜, 숫자, 장소)
   - 감정 과잉 표현 금지

[★★★ AI 티 나는 표현 절대 금지 ★★★]

다음 표현을 사용하면 실패입니다:
- "실수 1:", "실수 2:", "해결책:" 같은 나열
- "첫째,", "둘째,", "셋째," 같은 번호
- "중요합니다", "핵심입니다", "필수적입니다"
- "따라서", "그러므로", "결론적으로"
- "~라고 할 수 있습니다"
- "많은 분들이", "대부분의 사람들이"
- "~하는 것이 좋습니다"
- **굵은글씨**, *기울임*
- 1. 2. 3. 번호 목록
- 저는, (주어 뒤 쉼표)
- "포기하지 마세요", "도전해보세요" 같은 교훈

대신 이렇게 쓰세요:
- 자연스러운 문장 연결로 이야기 전개
- 구체적 사례와 숫자로 설명
- "저는 ~했습니다. 결과는 ~였습니다."

[올바른 예시]
"저는 처음 3개월 동안 단 한 건도 팔지 못했습니다. 매일 새벽까지 글을 썼는데 반응은 제로였죠. 뭐가 문제인지 정말 몰랐습니다.

그러다 한 가지를 바꿨습니다. 제목에서 '방법'이라는 단어를 빼고 구체적인 숫자를 넣은 거죠. '3개월 만에 월 300만원'처럼요. 그 다음 주에 첫 판매가 일어났습니다. 제목 하나 바꿨을 뿐인데 말이죠."

[분량]
1500~2000자 (공백 포함)

[미션]
'{subtopic_title}'의 본문만 작성하세요. 다른 내용 없이 오직 이 소제목만.
- 존댓말(합쇼체) 100% 유지
- AI 티 나는 표현 완전 배제
- 자연스러운 스토리텔링으로 몰입감 있게"""
    return ask_ai("베스트셀러 작가", prompt, temperature=0.75)


def refine_content(content, style="친근한"):
    style_guide = {
        "친근한": """친근한 스타일
- 합니다체(존댓말)로 통일
- 자신감 있는 단정 ("~입니다")
- 구체적 숫자와 팩트
- 독자에게 직접 말하듯""",
        
        "전문적": """전문가 스타일
- 합니다체(존댓말)로 통일
- 데이터와 출처 강조
- 논리적 전개
- 신뢰감 있는 톤""",
        
        "직설적": """직설 스타일
- 합니다체(존댓말)로 통일
- 핵심만 간결하게
- 군더더기 제로
- 팩트 중심""",
        
        "스토리텔링": """스토리 스타일
- 합니다체(존댓말)로 통일
- 구체적 장면 묘사
- 대화체 활용
- 감정선 살리기"""
    }
    
    prompt = f"""다음 글을 다듬어주세요.

[원본]
{content}

[★★★ 가장 중요: 문체 통일 ★★★]
1. 반드시 "합니다체(존댓말)"로 끝까지 통일
2. 반말("~다", "~했다")이 섞여 있으면 모두 존댓말로 수정
3. 모든 문장을 "~입니다", "~습니다", "~세요", "~죠"로 끝내기

[문단 구성 수정]
1. 한 문단은 반드시 2~4문장으로 구성
2. 절대로 한 문장만 띄어쓰기 하지 마세요
3. 관련된 문장들은 같은 문단에 묶으세요

[AI 티 제거 - 이런 표현 모두 수정]
- "실수 1:", "해결책:", "첫째, 둘째" → 자연스러운 문장으로 연결
- "중요합니다", "핵심입니다" 반복 → 다양한 표현으로 변경
- "따라서", "그러므로", "결론적으로" → 자연스러운 연결어로
- "~라고 할 수 있습니다" → 단정적 표현으로
- "많은 분들이", "누구나" → 구체적 표현으로

[추가 수정사항]
- 주어 뒤 쉼표 제거 (저는, → 저는)
- 마크다운 완전 제거 (**굵게**, *기울임*, 번호 매기기)
- 뻔한 비유 제거 또는 신선한 비유로 교체
- 교훈적 마무리 제거 ("포기하지 마세요" 등)

[목표 스타일]
{style_guide.get(style, style_guide["친근한"])}

[출력]
다듬어진 글만 출력. 설명 없이."""
    return ask_ai("에디터", prompt, temperature=0.7)

def check_quality(content):
    prompt = f"""당신은 "역행자", "부의 추월차선", "돈의 속성" 수준의 베스트셀러를 편집한 전문 편집자입니다.
다음 글이 베스트셀러 수준인지 냉정하게 평가해주세요.

[평가할 글]
{content[:4000]}

[평가 기준 - 베스트셀러 체크리스트]

1. 첫 문장 (10점)
   - 첫 문장이 독자의 뒤통수를 치는가?
   - 첫 문장만 읽고도 다음이 궁금한가?

2. 몰입도 (10점)
   - 중간에 멈추지 않고 끝까지 읽게 되는가?
   - 문장 리듬이 좋은가?

3. 공감력 (10점)
   - 독자가 "이건 내 얘기잖아"라고 느끼는가?
   - 타겟의 아픔을 정확히 건드리는가?

4. 구체성 (10점)
   - 추상적 조언 대신 구체적 장면/숫자가 있는가?
   - "열심히 했다" 대신 "새벽 4시에 일어났다" 수준인가?

5. AI 티 (10점, 감점 항목)
   - "~입니다" 반복, "따라서", "중요합니다" 등 AI 표현이 있는가?
   - 문장이 너무 균일하고 딱딱한가?

[출력 형식]

📊 종합 점수: __/50점

📌 첫 문장 평가: __/10점
- 현재 첫 문장: "[첫 문장 인용]"
- 평가: [좋은 점 또는 문제점]
- 개선안: "[더 좋은 첫 문장 제안]"

📌 몰입도 평가: __/10점
- [구체적 평가]

📌 공감력 평가: __/10점
- [구체적 평가]

📌 구체성 평가: __/10점
- [구체적 평가]

📌 AI 티 체크: __/10점
- 발견된 AI 표현: [있다면 나열]
- 개선이 필요한 문장: [3개까지]

✍️ 수정하면 좋을 문장 TOP 3
1. 원문: "..." → 수정안: "..."
2. 원문: "..." → 수정안: "..."
3. 원문: "..." → 수정안: "..."

💡 잘 쓴 문장 TOP 2
1. "[잘 쓴 문장]" - 좋은 이유
2. "[잘 쓴 문장]" - 좋은 이유

🎯 총평
[베스트셀러가 되기 위해 가장 중요한 개선점 1~2가지]"""
    return ask_ai("베스트셀러 편집자", prompt, temperature=0.6)

def generate_marketing_copy(title, subtitle, topic, persona):
    prompt = f"""당신은 크몽에서 전자책을 수천 권 판매한 탑셀러입니다.
당신의 상세페이지는 방문자의 15%가 구매하는 전설적인 전환율을 기록합니다.
당신의 카피는 읽는 순간 "이거 안 사면 손해"라는 느낌을 줍니다.

[상품 정보]
제목: {title}
부제: {subtitle}
주제: {topic}
타겟: {persona}

[미션]
이 전자책을 폭발적으로 팔기 위한 킬러 카피를 만들어주세요.

---

1. 크몽 상품 제목 (40자 이내)
   - 검색 키워드 포함 (SEO)
   - 구체적 결과/숫자 제시
   - 예시: "[PDF] 월 300벌게 해준 크몽 전자책 공식 | 실제 매출 인증"
   - 예시: "31개월 만에 10억 번 비밀 | 직장인 부업 전자책"

2. 상세페이지 헤드라인 3개
   - 스크롤을 멈추게 만드는 한 줄
   - 상식을 파괴하거나 충격을 줘야 함
   - 금지: "~하는 법", "~방법", "~가이드"
   - 예시: "월급만 믿다가는 평생 가난하다"
   - 예시: "나는 퇴사 3개월 만에 월급보다 더 벌었다"

3. 구매 유도 문구 (CTA) 3개
   - 긴급성 + FOMO(놓치면 후회) 자극
   - 구체적 숫자 활용
   - 예시: "이 가격은 100부 한정입니다"
   - 예시: "어제도 47명이 구매했습니다"
   - 예시: "지금 안 사면, 다음 달에는 2배입니다"

4. 인스타그램 홍보 문구
   - 첫 줄에서 스크롤 멈추게 (훅 필수)
   - 스토리텔링 요소 포함
   - 해시태그 5개 (검색량 높은 것)
   - 형식:
     [훅 - 첫 줄]
     
     [스토리 - 2~3줄]
     
     [CTA]
     
     #해시태그1 #해시태그2 ...

5. 블로그 포스팅 제목 3개
   - 검색 유입 + 클릭 유도
   - 궁금증 유발형
   - 예시: "크몽 전자책으로 월 500버는 사람들의 공통점 (실화)"
   - 예시: "직장인 부업 3개월 해본 후기 (feat. 월 수익 공개)"

---

모든 카피의 핵심 원칙:
- "이거 안 보면 나만 손해" 느낌
- 구체적 숫자로 신뢰감
- 호기심 자극 → 클릭 유도 → 구매 전환"""
    return ask_ai("크몽 탑셀러 마케터", prompt, temperature=0.85)


# ==========================================
# 🔧 글자 수 계산 헬퍼 함수 (통일된 계산 방식)
# ==========================================
def calculate_char_count(text):
    """순수 본문만으로 글자 수 계산 (공백, 줄바꿈 제외)"""
    if not text:
        return 0
    return len(text.replace('\n', '').replace(' ', ''))

def clean_content_for_display(content, subtopic_title=None, chapter_title=None):
    """본문에서 마크다운 기호, HTML 태그, 중복 제목, Unicode 제어 문자 제거"""
    if not content:
        return ""
    
    # 0. Unicode 제어 문자 제거 (RTL/LTR 마커, 제로폭 문자 등)
    # 이 문자들이 있으면 텍스트가 아랍어처럼 보일 수 있음
    unicode_control_chars = [
        '\u200e',  # LEFT-TO-RIGHT MARK
        '\u200f',  # RIGHT-TO-LEFT MARK
        '\u202a',  # LEFT-TO-RIGHT EMBEDDING
        '\u202b',  # RIGHT-TO-LEFT EMBEDDING
        '\u202c',  # POP DIRECTIONAL FORMATTING
        '\u202d',  # LEFT-TO-RIGHT OVERRIDE
        '\u202e',  # RIGHT-TO-LEFT OVERRIDE
        '\u2066',  # LEFT-TO-RIGHT ISOLATE
        '\u2067',  # RIGHT-TO-LEFT ISOLATE
        '\u2068',  # FIRST STRONG ISOLATE
        '\u2069',  # POP DIRECTIONAL ISOLATE
        '\u200b',  # ZERO WIDTH SPACE
        '\u200c',  # ZERO WIDTH NON-JOINER
        '\u200d',  # ZERO WIDTH JOINER
        '\ufeff',  # ZERO WIDTH NO-BREAK SPACE (BOM)
        '\u061c',  # ARABIC LETTER MARK
    ]
    for char in unicode_control_chars:
        content = content.replace(char, '')
    
    # 추가: 제어 문자 범위 제거 (U+0000 ~ U+001F, U+007F ~ U+009F)
    content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', content)
    
    # 1. HTML 태그 제거 (예: <div class="...">, </div>, <p>, </p> 등)
    content = re.sub(r'<[^>]+>', '', content)
    # HTML 엔티티 변환
    content = content.replace('&amp;', '&')
    content = content.replace('&lt;', '<')
    content = content.replace('&gt;', '>')
    content = content.replace('&quot;', '"')
    content = content.replace('&#39;', "'")
    content = content.replace('&nbsp;', ' ')
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for idx, line in enumerate(lines):
        stripped = line.strip()
        
        # 빈 줄은 그대로 유지 (단, 처음 3줄 이내에서는 건너뛰기)
        if not stripped:
            if idx > 3 or len(cleaned_lines) > 0:
                cleaned_lines.append(line)
            continue
        
        # 마크다운 헤더 제거 (##, ###)
        if stripped.startswith('#'):
            text_after = stripped.lstrip('#').strip()
            # 챕터 제목이나 소제목 관련이면 건너뛰기
            if chapter_title and (text_after in chapter_title or chapter_title in text_after):
                continue
            if subtopic_title and (text_after in subtopic_title or subtopic_title in text_after):
                continue
            # "챕터", "소제목" 키워드 포함하면 건너뛰기
            if '챕터' in text_after or '소제목' in text_after:
                continue
            continue  # 모든 마크다운 헤더 제거
        
        # "챕터 N:" 또는 "챕터N:" 형식 제거
        if stripped.startswith('챕터') and ':' in stripped[:15]:
            continue
        
        # "소제목:" 형식 제거
        if stripped.startswith('소제목') and ':' in stripped[:10]:
            continue
        
        # 처음 5줄 이내에서 소제목과 동일하거나 유사한 줄 제거
        if subtopic_title and idx < 5:
            clean_subtopic = subtopic_title.replace('**', '').strip()
            clean_stripped = stripped.replace('**', '').strip()
            if clean_stripped == clean_subtopic:
                continue
            # 소제목이 줄에 포함되어 있고 줄이 짧으면 제거
            if clean_subtopic in clean_stripped and len(clean_stripped) < len(clean_subtopic) + 20:
                continue
        
        # 챕터 제목과 동일한 줄 제거
        if chapter_title and idx < 5:
            clean_chapter = chapter_title.replace('**', '').strip()
            if clean_chapter in stripped or stripped in clean_chapter:
                continue
        
        cleaned_lines.append(line)
    
    # 결과 앞뒤 빈 줄 정리
    result = '\n'.join(cleaned_lines).strip()
    return result

def escape_html(text):
    """HTML 특수문자 이스케이프"""
    if not text:
        return ""
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))

def escape_rtf_unicode(text):
    """RTF용 유니코드 이스케이프 - 한글 등 비ASCII 문자 처리"""
    if not text:
        return ""
    result = []
    for char in text:
        code = ord(char)
        if code < 128:
            # ASCII 문자는 그대로
            if char == '\\':
                result.append('\\\\')
            elif char == '{':
                result.append('\\{')
            elif char == '}':
                result.append('\\}')
            elif char == '\n':
                result.append('\\line ')
            elif char == '\r':
                continue  # 캐리지 리턴 무시
            else:
                result.append(char)
        else:
            # 비ASCII 문자 (한글 등)는 유니코드 이스케이프
            # 음수 변환 처리 (signed 16-bit)
            if code > 32767:
                signed_code = code - 65536
            else:
                signed_code = code
            result.append(f'\\u{signed_code}?')
    return ''.join(result)

def get_all_content_text():
    """모든 챕터의 순수 본문 텍스트만 수집 (목차 순서 보장)"""
    pure_content = ""
    for ch in st.session_state.get('outline', []):
        if ch in st.session_state.get('chapters', {}):
            ch_data = st.session_state['chapters'][ch]
            if 'subtopic_data' in ch_data:
                subtopic_list = ch_data.get('subtopics', [])
                # 에필로그 등 소제목 없이 챕터 자체가 키인 경우
                if not subtopic_list and ch in ch_data['subtopic_data']:
                    subtopic_list = [ch]
                for st_name in subtopic_list:
                    st_data = ch_data['subtopic_data'].get(st_name, {})
                    if st_data.get('content'):
                        pure_content += st_data['content']
    return pure_content


# --- 메인 UI ---
st.markdown("""
<div class="hero-section">
    <div class="hero-label">CASHMAKER</div>
    <div class="hero-title">전자책 작성 프로그램</div>
    <div class="hero-subtitle">쉽고, 빠른 전자책 수익화</div>
</div>
""", unsafe_allow_html=True)

# 메인 탭
tabs = st.tabs([
    "① 주제 선정", 
    "② 타겟 & 컨셉", 
    "③ 목차 설계", 
    "④ 본문 작성", 
    "⑤ 문체 다듬기",
    "⑥ 최종 출력"
])

# === TAB 1: 주제 선정 ===
with tabs[0]:
    st.markdown("## 주제 선정 & 적합도 분석")
    
    # 빠른 시작 안내
    st.markdown("""
    <div class="quick-action-box">
        <p>💡 <strong>이미 주제가 있다면?</strong> 아래에 입력 후 바로 다음 탭으로 이동하세요!</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
        st.markdown("### 주제 입력")
        
        topic_input = st.text_input(
            "어떤 주제로 전자책을 쓰고 싶으세요?",
            value=st.session_state['topic'],
            placeholder="예: 크몽으로 월 500만원 벌기"
        )
        
        if topic_input != st.session_state['topic']:
            st.session_state['topic'] = topic_input
            st.session_state['topic_score'] = None
            st.session_state['score_details'] = None
        
        st.markdown("""
        <div class="info-card">
            <div class="info-card-title">좋은 주제의 조건</div>
            <p>• 내가 직접 경험하고 성과를 낸 것</p>
            <p>• 사람들이 돈 주고 배우고 싶어하는 것</p>
            <p>• 구체적인 결과를 약속할 수 있는 것</p>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("📊 적합도 분석하기 (선택)", key="analyze_btn"):
            if not topic_input:
                st.error("주제를 입력해주세요.")
            else:
                with st.spinner("분석 중..."):
                    result = analyze_topic_score(topic_input)
                    try:
                        json_match = re.search(r'\{[\s\S]*\}', result)
                        if json_match:
                            score_data = json.loads(json_match.group())
                            st.session_state['topic_score'] = score_data.get('total_score', 0)
                            st.session_state['topic_verdict'] = score_data.get('verdict', '분석 실패')
                            st.session_state['score_details'] = score_data
                    except:
                        st.error("분석 결과 파싱 오류. 다시 시도해주세요.")
    
    with col2:
        st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
        st.markdown("### 분석 결과")
        
        if st.session_state['topic_score'] is not None:
            score = st.session_state['topic_score']
            verdict = st.session_state['topic_verdict']
            details = st.session_state['score_details']
            
            verdict_class = "status-excellent" if verdict == "적합" else ("status-good" if verdict == "보통" else "status-warning")
            
            st.markdown(f"""
            <div class="score-card">
                <div class="score-number">{score}</div>
                <div class="score-label">종합 점수</div>
                <span class="status-badge {verdict_class}">{verdict}</span>
            </div>
            """, unsafe_allow_html=True)
            
            if details:
                st.markdown("#### 세부 점수")
                
                items = [
                    ("시장성", details.get('market', {}).get('score', 0), details.get('market', {}).get('reason', '')),
                    ("수익성", details.get('profit', {}).get('score', 0), details.get('profit', {}).get('reason', '')),
                    ("차별화", details.get('differentiation', {}).get('score', 0), details.get('differentiation', {}).get('reason', '')),
                    ("작성 난이도", details.get('difficulty', {}).get('score', 0), details.get('difficulty', {}).get('reason', '')),
                    ("지속성", details.get('sustainability', {}).get('score', 0), details.get('sustainability', {}).get('reason', '')),
                ]
                
                for name, score_val, reason in items:
                    st.markdown(f"""
                    <div class="score-item">
                        <span class="score-item-label">{name}</span>
                        <span class="score-item-value">{score_val}</span>
                    </div>
                    <p class="score-item-reason">{reason}</p>
                    """, unsafe_allow_html=True)
                
                st.markdown(f"""
                <div class="summary-box">
                    <p><strong>종합 의견</strong><br>{details.get('summary', '')}</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
                <p>분석은 선택사항입니다.</p>
                <p>주제만 입력해도 다음 단계로 진행 가능!</p>
            </div>
            """, unsafe_allow_html=True)

# === TAB 2: 타겟 & 컨셉 ===
with tabs[1]:
    st.markdown("## 타겟 설정 & 제목 생성")
    
    # 빠른 시작 안내
    if not st.session_state['topic']:
        st.info("💡 주제를 먼저 입력하면 더 정확한 결과를 얻을 수 있어요. 또는 여기서 바로 시작해도 됩니다!")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
        st.markdown("### 타겟 정의")
        
        # 주제가 없으면 여기서도 입력 가능
        if not st.session_state['topic']:
            topic_here = st.text_input(
                "주제 (여기서 입력 가능)",
                value=st.session_state['topic'],
                placeholder="예: 크몽으로 월 500만원 벌기",
                key="topic_tab2"
            )
            if topic_here:
                st.session_state['topic'] = topic_here
        
        persona = st.text_area(
            "누가 이 책을 읽나요?",
            value=st.session_state['target_persona'],
            placeholder="예: 30대 직장인, 퇴근 후 부업으로 월 100만원 추가 수입을 원하는 사람",
            height=100
        )
        st.session_state['target_persona'] = persona
        
        pain_points = st.text_area(
            "타겟의 가장 큰 고민은?",
            value=st.session_state['pain_points'],
            placeholder="예: 시간이 없다, 뭘 해야 할지 모르겠다, 시작이 두렵다",
            height=100
        )
        st.session_state['pain_points'] = pain_points
        
        st.markdown("---")
        
        st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
        st.markdown("### 한 줄 컨셉")
        
        if st.button("컨셉 생성하기", key="concept_btn"):
            if not st.session_state['topic'] or not persona:
                st.error("주제와 타겟을 먼저 입력해주세요.")
            else:
                with st.spinner("생성 중..."):
                    concept = generate_concept(
                        st.session_state['topic'],
                        persona,
                        pain_points
                    )
                    st.session_state['one_line_concept'] = concept
        
        if st.session_state['one_line_concept']:
            st.markdown(f"""
            <div class="info-card">
                {st.session_state['one_line_concept'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown('<p class="section-label">Step 03</p>', unsafe_allow_html=True)
        st.markdown("### 제목 생성")
        
        if st.button("제목 생성하기", key="title_btn"):
            if not st.session_state['topic']:
                st.error("주제를 먼저 입력해주세요.")
            else:
                with st.spinner("생성 중..."):
                    titles_result = generate_titles_advanced(
                        st.session_state['topic'],
                        st.session_state['target_persona'],
                        st.session_state['pain_points']
                    )
                    try:
                        json_match = re.search(r'\{[\s\S]*\}', titles_result)
                        if json_match:
                            st.session_state['generated_titles'] = json.loads(json_match.group())
                    except:
                        st.session_state['generated_titles'] = None
                        st.markdown(titles_result)
        
        if st.session_state.get('generated_titles'):
            titles_data = st.session_state['generated_titles']
            if 'titles' in titles_data:
                for i, t in enumerate(titles_data['titles'], 1):
                    st.markdown(f"""
                    <div class="title-card">
                        <div class="card-number">TITLE 0{i}</div>
                        <div class="main-title">{t.get('title', '')}</div>
                        <div class="sub-title">{t.get('subtitle', '')}</div>
                        <div class="reason">{t.get('why_works', '')}</div>
                    </div>
                    """, unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown('<p class="section-label">Step 04</p>', unsafe_allow_html=True)
        st.markdown("### 최종 선택")
        st.session_state['book_title'] = st.text_input("제목", value=st.session_state['book_title'], placeholder="최종 제목")
        st.session_state['subtitle'] = st.text_input("부제", value=st.session_state['subtitle'], placeholder="부제")

# === TAB 3: 목차 설계 ===
with tabs[2]:
    st.markdown("## 목차 설계")
    
    # 모드 선택
    st.markdown("### 🎯 작업 방식 선택")
    outline_mode = st.radio(
        "목차를 어떻게 만드시겠어요?",
        ["🤖 자동으로 목차 생성", "✍️ 내가 직접 입력"],
        horizontal=True,
        key="outline_mode_radio"
    )
    
    st.markdown("---")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        if outline_mode == "🤖 자동으로 목차 생성":
            st.markdown('<p class="section-label">자동 목차 생성</p>', unsafe_allow_html=True)
            st.markdown("### 목차를 자동으로 설계합니다")
            
            # 주제가 없으면 여기서도 입력 가능
            if not st.session_state['topic']:
                st.warning("💡 주제를 먼저 입력해주세요")
                topic_here = st.text_input(
                    "주제",
                    value=st.session_state['topic'],
                    placeholder="예: 크몽으로 월 500만원 벌기",
                    key="topic_tab3"
                )
                if topic_here:
                    st.session_state['topic'] = topic_here
            
            if st.button("🚀 목차 생성하기", key="outline_btn"):
                if not st.session_state['topic']:
                    st.error("주제를 먼저 입력해주세요.")
                else:
                    with st.spinner("설계 중..."):
                        outline_text = generate_outline(
                            st.session_state['topic'],
                            st.session_state['target_persona'],
                            st.session_state['pain_points']
                        )
                        
                        # 챕터와 소제목 파싱
                        lines = outline_text.split('\n')
                        chapters = []
                        current_chapter = None
                        chapter_subtopics = {}
                        
                        for line in lines:
                            line = line.strip()
                            if not line or line == '...':
                                continue
                            
                            # 챕터 판별 함수
                            def is_chapter_line_ai(text):
                                text_clean = text.lstrip('#').strip()
                                text_lower = text_clean.lower()
                                # 키워드로 시작
                                if any(text_lower.startswith(kw) for kw in ['챕터', 'chapter', '에필로그', '프롤로그', '서문', '부록']):
                                    return True
                                # X부. X부: X장. X장: 형태
                                if len(text_clean) > 1 and text_clean[0].isdigit():
                                    rest = text_clean[1:].lstrip('0123456789')
                                    if rest and (rest[0] in '부장.:'):
                                        return True
                                return False
                            
                            # 소제목 판별 함수
                            def is_subtopic_line_ai(text):
                                if text[0] in '-·•':
                                    return True
                                # 숫자) 형태
                                if len(text) > 1 and text[0].isdigit():
                                    for i, char in enumerate(text):
                                        if char == ')':
                                            return True
                                        if not char.isdigit():
                                            break
                                return False
                            
                            # 챕터인지 확인
                            if is_chapter_line_ai(line):
                                chapter_name = line.lstrip('#').strip()
                                current_chapter = chapter_name
                                chapters.append(current_chapter)
                                chapter_subtopics[current_chapter] = []
                            elif current_chapter:
                                if is_subtopic_line_ai(line):
                                    subtopic = line.lstrip('-·• ')
                                    subtopic = re.sub(r'^\d+\)\s*', '', subtopic)
                                    if subtopic:
                                        chapter_subtopics[current_chapter].append(subtopic)
                        
                        # 저장
                        st.session_state['outline'] = chapters
                        st.session_state['full_outline'] = outline_text
                        
                        # 챕터별 데이터 생성
                        for ch in chapters:
                            subtopics = chapter_subtopics.get(ch, [])
                            st.session_state['chapters'][ch] = {
                                'subtopics': subtopics,
                                'subtopic_data': {st: {'questions': [], 'answers': [], 'content': ''} for st in subtopics}
                            }
                        
                        # 결과 표시
                        total_subtopics = sum(len(chapter_subtopics.get(ch, [])) for ch in chapters)
                        st.success(f"✅ {len(chapters)}개 챕터, {total_subtopics}개 소제목 생성됨!")
                        st.rerun()
            
            if 'full_outline' in st.session_state and st.session_state['full_outline']:
                # sync_full_outline으로 생성된 최신 목차 표시 (읽기 전용)
                st.markdown("**📋 현재 목차**")
                st.code(st.session_state['full_outline'], language=None)
        
        else:  # 직접 입력 모드
            st.markdown('<p class="section-label">직접 입력</p>', unsafe_allow_html=True)
            st.markdown("### 목차를 직접 입력하세요")
            
            st.markdown("""
            <div class="info-card">
                <div class="info-card-title">📌 입력 형식 예시</div>
                <p><b>1부. [KODE 1] 크몽에서 월 1,000 창출하라</b></p>
                <p style="margin-left: 20px;">1) 왕따가 억대 사업가가 될 수 있었던 이유</p>
                <p style="margin-left: 20px;">2) 리버스 엔지니어링 4단계</p>
                <p><b>2부. [KODE 2] 선수익 월배당 시스템</b></p>
                <p style="margin-left: 20px;">3) 자본주의를 해킹하는 법</p>
                <p style="margin-left: 20px;">4) 듀얼 엔진 이론</p>
            </div>
            """, unsafe_allow_html=True)
            
            # 기존 목차가 있으면 소제목 포함해서 표시
            existing_outline = ""
            if st.session_state['outline']:
                for ch in st.session_state['outline']:
                    existing_outline += f"{ch}\n"
                    if ch in st.session_state['chapters']:
                        for i, st_name in enumerate(st.session_state['chapters'][ch].get('subtopics', []), 1):
                            existing_outline += f"{i}) {st_name}\n"
            
            manual_outline = st.text_area(
                "목차 입력 (챕터와 소제목)",
                value=existing_outline,
                height=350,
                placeholder="1부. 첫 번째 챕터 제목\n1) 소제목 1\n2) 소제목 2\n2부. 두 번째 챕터 제목\n3) 소제목 3\n...",
                key="manual_outline_input"
            )
            
            if st.button("✅ 목차 저장하기", key="save_manual_outline"):
                if manual_outline.strip():
                    # 챕터와 소제목 파싱
                    lines = manual_outline.strip().split('\n')
                    chapters = []
                    current_chapter = None
                    chapter_subtopics = {}
                    
                    for line in lines:
                        line = line.strip()
                        if not line or line == '...':
                            continue
                        
                        # 챕터 판별: X부., X장., X., 챕터X, 에필로그 등
                        is_chapter = False
                        text_lower = line.lower()
                        
                        # 키워드로 시작하는 경우
                        if any(text_lower.startswith(kw) for kw in ['챕터', 'chapter', '에필로그', '프롤로그', '서문', '부록']):
                            is_chapter = True
                        # 숫자로 시작하는 경우
                        elif len(line) > 2 and line[0].isdigit():
                            # 숫자 부분 추출
                            num_end = 1
                            while num_end < len(line) and line[num_end].isdigit():
                                num_end += 1
                            rest = line[num_end:]
                            
                            if rest:
                                # "부", "장" 다음에 오면 챕터
                                if rest[0] in '부장':
                                    is_chapter = True
                                # "숫자. " 형태 (점 다음에 공백이나 문자가 오면) - 소제목 아님
                                elif rest[0] == '.' and len(rest) > 1 and rest[1] != ')':
                                    is_chapter = True
                                # "숫자: " 형태
                                elif rest[0] == ':':
                                    is_chapter = True
                        
                        # 소제목 판별: X), -, ·, • 로 시작
                        is_subtopic = False
                        if not is_chapter and current_chapter:
                            # -, ·, • 로 시작
                            if line[0] in '-·•':
                                is_subtopic = True
                            # 숫자) 형태
                            elif line[0].isdigit():
                                num_end = 1
                                while num_end < len(line) and line[num_end].isdigit():
                                    num_end += 1
                                if num_end < len(line) and line[num_end] == ')':
                                    is_subtopic = True
                        
                        # 처리
                        if is_chapter:
                            current_chapter = line
                            chapters.append(current_chapter)
                            chapter_subtopics[current_chapter] = []
                        elif current_chapter:
                            # 소제목으로 처리
                            subtopic = line
                            # 접두사 제거
                            if line[0] in '-·•':
                                subtopic = line.lstrip('-·• ').strip()
                            elif is_subtopic:
                                # 숫자) 제거
                                subtopic = re.sub(r'^\d+\)\s*', '', line).strip()
                            
                            if subtopic and len(subtopic) > 2:
                                chapter_subtopics[current_chapter].append(subtopic)
                    
                    # 저장
                    st.session_state['outline'] = chapters
                    st.session_state['full_outline'] = manual_outline
                    
                    # 챕터별 데이터 생성
                    for ch in chapters:
                        subtopics = chapter_subtopics.get(ch, [])
                        st.session_state['chapters'][ch] = {
                            'subtopics': subtopics,
                            'subtopic_data': {st_name: {'questions': [], 'answers': [], 'content': ''} for st_name in subtopics}
                        }
                    
                    trigger_auto_save()
                    
                    # 결과 표시
                    total_subtopics = sum(len(chapter_subtopics.get(ch, [])) for ch in chapters)
                    st.success(f"✅ {len(chapters)}개 챕터, {total_subtopics}개 소제목 저장됨!")
                    st.rerun()
                else:
                    st.error("목차를 입력해주세요.")
    
    with col2:
        st.markdown('<p class="section-label">목차 관리</p>', unsafe_allow_html=True)
        st.markdown("### 📋 현재 목차")
        
        if st.session_state['outline']:
            # 챕터별 표시 및 관리
            for i, chapter in enumerate(st.session_state['outline']):
                subtopic_count = 0
                if chapter in st.session_state['chapters']:
                    subtopic_count = len(st.session_state['chapters'][chapter].get('subtopics', []))
                
                # 챕터 헤더
                with st.expander(f"**{chapter}** ({subtopic_count}개 소제목)", expanded=False):
                    # 챕터 제목 편집
                    col_edit, col_actions = st.columns([3, 2])
                    with col_edit:
                        new_title = st.text_input(
                            "챕터 제목",
                            value=chapter,
                            key=f"edit_chapter_{i}",
                            label_visibility="collapsed"
                        )
                    with col_actions:
                        col_regen, col_del = st.columns(2)
                        with col_regen:
                            if st.button("🔄 재생성", key=f"regen_chapter_{i}", help="이 챕터만 새로 생성"):
                                with st.spinner("챕터 재생성 중..."):
                                    new_chapter_text = regenerate_chapter_outline(
                                        i + 1,
                                        st.session_state['topic'],
                                        st.session_state['target_persona'],
                                        st.session_state['outline']
                                    )
                                    # 파싱
                                    lines = new_chapter_text.split('\n')
                                    new_chapter_title = None
                                    new_subtopics = []
                                    for line in lines:
                                        line = line.strip()
                                        if line.startswith('##'):
                                            new_chapter_title = line.lstrip('#').strip()
                                        elif line.startswith('-'):
                                            st_name = line.lstrip('- ').strip()
                                            if st_name:
                                                new_subtopics.append(st_name)
                                    
                                    if new_chapter_title:
                                        old_chapter = st.session_state['outline'][i]
                                        st.session_state['outline'][i] = new_chapter_title
                                        # 기존 챕터 데이터 삭제
                                        if old_chapter in st.session_state['chapters']:
                                            del st.session_state['chapters'][old_chapter]
                                        # 새 챕터 데이터 생성
                                        st.session_state['chapters'][new_chapter_title] = {
                                            'subtopics': new_subtopics,
                                            'subtopic_data': {st: {'questions': [], 'answers': [], 'content': ''} for st in new_subtopics}
                                        }
                                        trigger_auto_save()
                                        st.rerun()
                        with col_del:
                            if st.button("🗑️", key=f"del_chapter_{i}", help="삭제"):
                                old_chapter = st.session_state['outline'].pop(i)
                                if old_chapter in st.session_state['chapters']:
                                    del st.session_state['chapters'][old_chapter]
                                trigger_auto_save()
                                st.rerun()
                    
                    # 챕터 제목 변경 저장
                    if new_title != chapter and new_title.strip():
                        if st.button("💾 제목 저장", key=f"save_chapter_title_{i}"):
                            st.session_state['outline'][i] = new_title
                            if chapter in st.session_state['chapters']:
                                st.session_state['chapters'][new_title] = st.session_state['chapters'].pop(chapter)
                            trigger_auto_save()
                            st.rerun()
                    
                    st.markdown("---")
                    st.markdown("**📝 소제목 관리**")
                    
                    # 소제목 목록
                    if chapter in st.session_state['chapters']:
                        subtopics = st.session_state['chapters'][chapter].get('subtopics', [])
                        
                        for j, st_name in enumerate(subtopics):
                            col_st, col_st_actions = st.columns([3, 2])
                            with col_st:
                                new_st = st.text_input(
                                    f"소제목 {j+1}",
                                    value=st_name,
                                    key=f"edit_st_{i}_{j}",
                                    label_visibility="collapsed"
                                )
                            with col_st_actions:
                                col_st_save, col_st_regen, col_st_del = st.columns(3)
                                with col_st_save:
                                    if new_st != st_name and new_st.strip():
                                        if st.button("💾", key=f"save_st_{i}_{j}", help="저장"):
                                            # 소제목 변경
                                            st.session_state['chapters'][chapter]['subtopics'][j] = new_st
                                            # subtopic_data 키도 변경
                                            if st_name in st.session_state['chapters'][chapter]['subtopic_data']:
                                                st.session_state['chapters'][chapter]['subtopic_data'][new_st] = st.session_state['chapters'][chapter]['subtopic_data'].pop(st_name)
                                            trigger_auto_save()
                                            st.rerun()
                                with col_st_regen:
                                    if st.button("🔄", key=f"regen_st_{i}_{j}", help="재생성"):
                                        with st.spinner("소제목 재생성 중..."):
                                            new_st_title = regenerate_single_subtopic(
                                                chapter,
                                                j + 1,
                                                st.session_state['topic'],
                                                subtopics
                                            )
                                            if new_st_title:
                                                old_st = st.session_state['chapters'][chapter]['subtopics'][j]
                                                st.session_state['chapters'][chapter]['subtopics'][j] = new_st_title
                                                # subtopic_data 키도 변경
                                                if old_st in st.session_state['chapters'][chapter]['subtopic_data']:
                                                    st.session_state['chapters'][chapter]['subtopic_data'][new_st_title] = st.session_state['chapters'][chapter]['subtopic_data'].pop(old_st)
                                                else:
                                                    st.session_state['chapters'][chapter]['subtopic_data'][new_st_title] = {'questions': [], 'answers': [], 'content': ''}
                                                trigger_auto_save()
                                                st.rerun()
                                with col_st_del:
                                    if st.button("🗑️", key=f"del_st_{i}_{j}", help="삭제"):
                                        removed_st = st.session_state['chapters'][chapter]['subtopics'].pop(j)
                                        if removed_st in st.session_state['chapters'][chapter]['subtopic_data']:
                                            del st.session_state['chapters'][chapter]['subtopic_data'][removed_st]
                                        trigger_auto_save()
                                        st.rerun()
                        
                        # 소제목 추가
                        st.markdown("---")
                        col_add_st, col_add_btn = st.columns([3, 1])
                        with col_add_st:
                            new_st_input = st.text_input(
                                "새 소제목",
                                placeholder="새 소제목을 입력하세요",
                                key=f"new_st_input_{i}",
                                label_visibility="collapsed"
                            )
                        with col_add_btn:
                            if st.button("➕ 추가", key=f"add_st_{i}"):
                                if new_st_input.strip():
                                    st.session_state['chapters'][chapter]['subtopics'].append(new_st_input.strip())
                                    st.session_state['chapters'][chapter]['subtopic_data'][new_st_input.strip()] = {
                                        'questions': [], 'answers': [], 'content': ''
                                    }
                                    trigger_auto_save()
                                    st.rerun()
                    else:
                        st.info("소제목이 없습니다. 아래에서 추가하세요.")
            
            st.markdown("---")
            
            # 챕터 추가 버튼
            if st.button("➕ 새 챕터 추가", key="add_chapter"):
                new_ch_name = f"챕터{len(st.session_state['outline'])+1}: 새 챕터"
                st.session_state['outline'].append(new_ch_name)
                st.session_state['chapters'][new_ch_name] = {
                    'subtopics': [],
                    'subtopic_data': {}
                }
                trigger_auto_save()
                st.rerun()
            
        else:
            st.markdown("""
            <div class="empty-state">
                <p>왼쪽에서 목차를 생성하거나 직접 입력하세요</p>
            </div>
            """, unsafe_allow_html=True)

# === TAB 4: 본문 작성 ===
with tabs[3]:
    st.markdown("## 본문 작성")
    
    # 목차가 없는 경우
    if not st.session_state['outline']:
        st.warning("⚠️ 먼저 '③ 목차 설계' 탭에서 목차를 작성해주세요.")
        st.stop()
    
    # 챕터만 필터링 (소제목 제외)
    chapter_list = []
    for item in st.session_state['outline']:
        # 소제목인지 확인 (-, ·, • 로 시작하면 소제목)
        item_stripped = item.strip()
        if not item_stripped.startswith('-') and not item_stripped.startswith('·') and not item_stripped.startswith('•'):
            chapter_list.append(item)
    
    if not chapter_list:
        st.warning("⚠️ 챕터가 없습니다. 목차를 다시 확인해주세요.")
        st.stop()
    
    # 챕터 선택
    selected_chapter = st.selectbox(
        "📚 챕터 선택",
        chapter_list,
        key="chapter_select_main"
    )
    
    # 챕터 데이터 초기화
    if selected_chapter not in st.session_state['chapters']:
        st.session_state['chapters'][selected_chapter] = {
            'subtopics': [],
            'subtopic_data': {}
        }
    
    chapter_data = st.session_state['chapters'][selected_chapter]
    
    if 'subtopics' not in chapter_data:
        chapter_data['subtopics'] = []
    if 'subtopic_data' not in chapter_data:
        chapter_data['subtopic_data'] = {}
    
    # 기존 소제목 데이터 초기화
    for st_name in chapter_data['subtopics']:
        if st_name not in chapter_data['subtopic_data']:
            chapter_data['subtopic_data'][st_name] = {'questions': [], 'answers': [], 'content': ''}
    
    st.markdown("---")
    
    # 📋 현재 챕터의 소제목 전체 보기 (확장 가능)
    with st.expander(f"📋 '{selected_chapter}' 소제목 전체 보기 ({len(chapter_data.get('subtopics', []))}개)", expanded=False):
        if chapter_data.get('subtopics'):
            for j, st_name in enumerate(chapter_data['subtopics']):
                # 작성 완료 여부 확인
                has_content = bool(chapter_data['subtopic_data'].get(st_name, {}).get('content', '').strip())
                status_icon = "✅" if has_content else "⬜"
                
                col_st_view, col_st_edit, col_st_regen = st.columns([4, 1, 1])
                with col_st_view:
                    new_st_name = st.text_input(
                        f"{status_icon} {j+1}",
                        value=st_name,
                        key=f"view_st_tab4_{j}",
                        label_visibility="collapsed"
                    )
                with col_st_edit:
                    if new_st_name != st_name and new_st_name.strip():
                        if st.button("💾", key=f"save_st_tab4_{j}", help="저장"):
                            chapter_data['subtopics'][j] = new_st_name
                            if st_name in chapter_data['subtopic_data']:
                                chapter_data['subtopic_data'][new_st_name] = chapter_data['subtopic_data'].pop(st_name)
                            trigger_auto_save()
                            st.rerun()
                with col_st_regen:
                    if st.button("🔄", key=f"regen_st_tab4_{j}", help="이 소제목만 재생성"):
                        with st.spinner("재생성 중..."):
                            new_title = regenerate_single_subtopic(
                                selected_chapter,
                                j + 1,
                                st.session_state['topic'],
                                chapter_data['subtopics']
                            )
                            if new_title:
                                old_st = chapter_data['subtopics'][j]
                                chapter_data['subtopics'][j] = new_title
                                if old_st in chapter_data['subtopic_data']:
                                    chapter_data['subtopic_data'][new_title] = chapter_data['subtopic_data'].pop(old_st)
                                else:
                                    chapter_data['subtopic_data'][new_title] = {'questions': [], 'answers': [], 'content': ''}
                                trigger_auto_save()
                                st.rerun()
            
            # 새 소제목 추가
            st.markdown("---")
            col_new_st, col_new_btn = st.columns([4, 1])
            with col_new_st:
                new_st_input = st.text_input(
                    "새 소제목 추가",
                    placeholder="새 소제목을 입력하세요",
                    key="new_st_input_tab4",
                    label_visibility="collapsed"
                )
            with col_new_btn:
                if st.button("➕", key="add_st_tab4", help="추가"):
                    if new_st_input.strip():
                        chapter_data['subtopics'].append(new_st_input.strip())
                        chapter_data['subtopic_data'][new_st_input.strip()] = {
                            'questions': [], 'answers': [], 'content': ''
                        }
                        trigger_auto_save()
                        st.rerun()
        else:
            st.info("소제목이 없습니다.")
    
    st.markdown("---")
    
    # ====== 소제목이 있는 경우: 바로 본문 작업 ======
    if chapter_data['subtopics']:
        st.markdown("### ✍️ 소제목 선택 → 본문 작성")
        
        # 소제목 선택 (selectbox)
        selected_subtopic = st.selectbox(
            "작성할 소제목",
            chapter_data['subtopics'],
            key="subtopic_select_main",
            format_func=lambda x: f"{'✅' if chapter_data['subtopic_data'].get(x, {}).get('content') else '⬜'} {x}"
        )
        
        # 진행 상황 표시
        completed = sum(1 for s in chapter_data['subtopics'] if chapter_data['subtopic_data'].get(s, {}).get('content'))
        total = len(chapter_data['subtopics'])
        st.progress(completed / total if total > 0 else 0)
        st.caption(f"진행: {completed}/{total} 완료")
        
        st.markdown("---")
        
        # 본문 작성 영역
        if selected_subtopic:
            if selected_subtopic not in chapter_data['subtopic_data']:
                chapter_data['subtopic_data'][selected_subtopic] = {'questions': [], 'answers': [], 'content': ''}
            
            subtopic_data = chapter_data['subtopic_data'][selected_subtopic]
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
                st.markdown(f"### 🎤 인터뷰: {selected_subtopic}")
                
                if st.button("🎤 질문 생성하기", key="gen_questions_main"):
                    with st.spinner("질문 생성 중..."):
                        questions_text = generate_interview_questions(
                            selected_subtopic, 
                            selected_chapter, 
                            st.session_state['topic']
                        )
                        questions = re.findall(r'Q\d+:\s*(.+)', questions_text)
                        if not questions:
                            questions = [q.strip() for q in questions_text.split('\n') if q.strip() and '?' in q][:3]
                        subtopic_data['questions'] = questions
                        subtopic_data['answers'] = [''] * len(questions)
                        st.rerun()
                
                if subtopic_data['questions']:
                    for i, q in enumerate(subtopic_data['questions']):
                        st.markdown(f"**Q{i+1}.** {q}")
                        if i >= len(subtopic_data['answers']):
                            subtopic_data['answers'].append('')
                        subtopic_data['answers'][i] = st.text_area(
                            f"A{i+1}",
                            value=subtopic_data['answers'][i],
                            key=f"answer_main_{selected_chapter}_{selected_subtopic}_{i}",
                            height=80,
                            label_visibility="collapsed"
                        )
                else:
                    st.info("👆 '질문 생성하기' 버튼을 눌러 인터뷰를 시작하세요.")
            
            with col2:
                st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
                st.markdown(f"### 📝 본문: {selected_subtopic}")
                
                # 본문 생성 조건 체크
                has_answers = subtopic_data.get('questions') and any(a.strip() for a in subtopic_data.get('answers', []))
                
                # 🔧 수정: 위젯 키 정의
                content_widget_key = f"content_main_{selected_chapter}_{selected_subtopic}"
                
                if has_answers:
                    if st.button("✨ 본문 생성하기", key="gen_content_main"):
                        with st.spinner("집필 중... (30초~1분)"):
                            content = generate_subtopic_content(
                                selected_subtopic,
                                selected_chapter,
                                subtopic_data['questions'],
                                subtopic_data['answers'],
                                st.session_state['topic'],
                                st.session_state['target_persona']
                            )
                            # 🔧 수정: 데이터와 위젯 상태 모두 업데이트
                            st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic]['content'] = content
                            st.session_state[content_widget_key] = content  # 위젯 키도 업데이트
                            trigger_auto_save()
                            st.rerun()
                else:
                    st.info("👈 먼저 인터뷰 질문에 답변해주세요.")
                
                # 🔧 수정: 현재 저장된 본문 가져오기 (위젯 상태 우선, 없으면 데이터에서)
                # 🔧 수정: 위젯 렌더링 전에 session_state 초기화 (value 사용 안 함)
                stored_content = st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic].get('content', '')
                
                # 🔧 핵심 수정: 항상 저장된 데이터와 위젯 키를 동기화
                # 위젯 키가 없거나, 현재 선택된 소제목이 변경되었을 때 재설정
                current_selection_key = f"_last_selected_{selected_chapter}"
                last_selected = st.session_state.get(current_selection_key, None)
                
                if last_selected != selected_subtopic:
                    # 소제목이 변경됨 - 위젯 키를 저장된 데이터로 재설정
                    st.session_state[content_widget_key] = stored_content
                    st.session_state[current_selection_key] = selected_subtopic
                elif content_widget_key not in st.session_state:
                    st.session_state[content_widget_key] = stored_content
                
                # 본문 표시 및 편집
                edited_content = st.text_area(
                    "본문 내용",
                    height=400,
                    key=content_widget_key,
                    label_visibility="collapsed"
                )
                
                # 편집된 내용 저장 (위젯에서 직접 가져옴)
                if content_widget_key in st.session_state:
                    st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic]['content'] = st.session_state[content_widget_key]
                
                # 🔧 수정: 글자 수 계산 - 저장된 데이터 기준
                final_content = st.session_state['chapters'][selected_chapter]['subtopic_data'][selected_subtopic].get('content', '')
                if final_content:
                    char_count = calculate_char_count(final_content)
                    st.caption(f"📊 {char_count:,}자")
                    st.success(f"✅ '{selected_subtopic}' 본문 작성 완료!")
        
        # 소제목 편집 (접혀있는 상태)
        with st.expander("⚙️ 소제목 편집/추가", expanded=False):
            st.markdown("#### 소제목 관리")
            
            col_gen, col_add = st.columns(2)
            
            with col_gen:
                num_subtopics = st.number_input(
                    "생성할 개수",
                    min_value=1,
                    max_value=10,
                    value=3,
                    key="num_subtopics_gen_exp"
                )
                if st.button("✨ 소제목 자동 생성", key="gen_subtopics_exp"):
                    with st.spinner("생성 중..."):
                        subtopics_text = generate_subtopics(
                            selected_chapter,
                            st.session_state['topic'],
                            st.session_state['target_persona'],
                            num_subtopics
                        )
                        new_subtopics = []
                        for line in subtopics_text.split('\n'):
                            line = line.strip()
                            if line and (line[0].isdigit() or line.startswith('-')):
                                cleaned = re.sub(r'^[\d\.\-\s]+', '', line).strip()
                                if cleaned:
                                    new_subtopics.append(cleaned)
                        
                        if new_subtopics:
                            chapter_data['subtopics'] = new_subtopics[:num_subtopics]
                            for st_name in new_subtopics[:num_subtopics]:
                                if st_name not in chapter_data['subtopic_data']:
                                    chapter_data['subtopic_data'][st_name] = {'questions': [], 'answers': [], 'content': ''}
                            st.success(f"✅ {len(new_subtopics[:num_subtopics])}개 생성됨!")
                            st.rerun()
            
            with col_add:
                new_name = st.text_input("새 소제목", placeholder="직접 입력", key="new_subtopic_exp")
                if st.button("➕ 추가", key="add_subtopic_exp"):
                    if new_name.strip() and new_name not in chapter_data['subtopics']:
                        chapter_data['subtopics'].append(new_name)
                        chapter_data['subtopic_data'][new_name] = {'questions': [], 'answers': [], 'content': ''}
                        st.rerun()
            
            # 현재 소제목 목록
            st.markdown("**현재 소제목:**")
            for i, st_name in enumerate(chapter_data['subtopics']):
                col_n, col_del = st.columns([5, 1])
                with col_n:
                    st.write(f"{i+1}. {st_name}")
                with col_del:
                    if st.button("🗑️", key=f"del_st_exp_{i}"):
                        chapter_data['subtopics'].remove(st_name)
                        if st_name in chapter_data['subtopic_data']:
                            del chapter_data['subtopic_data'][st_name]
                        st.rerun()
    
    # ====== 소제목이 없는 경우 (에필로그/프롤로그 등): 바로 본문 작성 ======
    else:
        # 에필로그, 프롤로그 등인지 확인
        is_special_chapter = any(kw in selected_chapter.lower() for kw in ['에필로그', '프롤로그', '서문', '부록', 'epilogue', 'prologue'])
        
        if is_special_chapter:
            st.info(f"📝 '{selected_chapter}'는 소제목 없이 바로 본문을 작성합니다.")
            
            # 챕터 자체를 소제목처럼 사용
            chapter_as_subtopic = selected_chapter
            if chapter_as_subtopic not in chapter_data['subtopic_data']:
                chapter_data['subtopic_data'][chapter_as_subtopic] = {'questions': [], 'answers': [], 'content': ''}
            
            subtopic_data = chapter_data['subtopic_data'][chapter_as_subtopic]
            
            # 🔧 수정: 위젯 키 정의
            content_widget_key_special = f"content_special_{selected_chapter}"
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown('<p class="section-label">Step 01</p>', unsafe_allow_html=True)
                st.markdown(f"### 🎤 인터뷰: {selected_chapter}")
                
                if st.button("🎤 질문 생성하기", key="gen_questions_special"):
                    with st.spinner("질문 생성 중..."):
                        questions_text = generate_interview_questions(
                            selected_chapter, 
                            selected_chapter, 
                            st.session_state['topic']
                        )
                        questions = re.findall(r'Q\d+:\s*(.+)', questions_text)
                        if not questions:
                            questions = [q.strip() for q in questions_text.split('\n') if q.strip() and '?' in q][:3]
                        subtopic_data['questions'] = questions
                        subtopic_data['answers'] = [''] * len(questions)
                        st.rerun()
                
                if subtopic_data['questions']:
                    for i, q in enumerate(subtopic_data['questions']):
                        st.markdown(f"**Q{i+1}.** {q}")
                        if i >= len(subtopic_data['answers']):
                            subtopic_data['answers'].append('')
                        subtopic_data['answers'][i] = st.text_area(
                            f"A{i+1}",
                            value=subtopic_data['answers'][i],
                            key=f"answer_special_{i}",
                            height=80,
                            label_visibility="collapsed"
                        )
                else:
                    st.info("👆 '질문 생성하기' 버튼을 눌러 인터뷰를 시작하세요.")
            
            with col2:
                st.markdown('<p class="section-label">Step 02</p>', unsafe_allow_html=True)
                st.markdown(f"### 📝 본문: {selected_chapter}")
                
                has_answers = subtopic_data.get('questions') and any(a.strip() for a in subtopic_data.get('answers', []))
                
                if has_answers:
                    if st.button("✨ 본문 생성하기", key="gen_content_special"):
                        with st.spinner("집필 중... (30초~1분)"):
                            content = generate_subtopic_content(
                                selected_chapter,
                                selected_chapter,
                                subtopic_data['questions'],
                                subtopic_data['answers'],
                                st.session_state['topic'],
                                st.session_state['target_persona']
                            )
                            # 🔧 수정: 데이터와 위젯 상태 모두 업데이트
                            st.session_state['chapters'][selected_chapter]['subtopic_data'][chapter_as_subtopic]['content'] = content
                            st.session_state[content_widget_key_special] = content
                            trigger_auto_save()
                            st.rerun()
                else:
                    st.info("👈 먼저 인터뷰 질문에 답변해주세요.")
                
                # 🔧 수정: 위젯 렌더링 전에 session_state 초기화 (value 사용 안 함)
                stored_content = st.session_state['chapters'][selected_chapter]['subtopic_data'].get(chapter_as_subtopic, {}).get('content', '')
                if content_widget_key_special not in st.session_state:
                    st.session_state[content_widget_key_special] = stored_content
                
                edited_content = st.text_area(
                    "본문 내용",
                    height=400,
                    key=content_widget_key_special,
                    label_visibility="collapsed"
                )
                
                # 편집된 내용 저장
                if content_widget_key_special in st.session_state:
                    st.session_state['chapters'][selected_chapter]['subtopic_data'][chapter_as_subtopic]['content'] = st.session_state[content_widget_key_special]
                
                final_content = st.session_state['chapters'][selected_chapter]['subtopic_data'].get(chapter_as_subtopic, {}).get('content', '')
                if final_content:
                    char_count = calculate_char_count(final_content)
                    st.caption(f"📊 {char_count:,}자")
                    st.success(f"✅ '{selected_chapter}' 본문 작성 완료!")
        
        else:
            st.warning("⚠️ 이 챕터에 소제목이 없습니다. 아래에서 소제목을 생성하거나 추가해주세요.")
            
            st.markdown("### 📝 소제목 생성")
            
            col_gen, col_add = st.columns(2)
            
            with col_gen:
                st.markdown("**자동 생성**")
                num_subtopics = st.number_input(
                    "생성할 개수",
                    min_value=1,
                    max_value=10,
                    value=3,
                    key="num_subtopics_gen_empty"
                )
                if st.button("✨ 소제목 자동 생성", key="gen_subtopics_empty"):
                    with st.spinner("베스트셀러급 소제목 생성 중..."):
                        subtopics_text = generate_subtopics(
                            selected_chapter,
                            st.session_state['topic'],
                            st.session_state['target_persona'],
                            num_subtopics
                        )
                        new_subtopics = []
                        for line in subtopics_text.split('\n'):
                            line = line.strip()
                            if line and (line[0].isdigit() or line.startswith('-')):
                                cleaned = re.sub(r'^[\d\.\-\s]+', '', line).strip()
                                if cleaned:
                                    new_subtopics.append(cleaned)
                        
                        if new_subtopics:
                            chapter_data['subtopics'] = new_subtopics[:num_subtopics]
                            for st_name in new_subtopics[:num_subtopics]:
                                chapter_data['subtopic_data'][st_name] = {'questions': [], 'answers': [], 'content': ''}
                            st.success(f"✅ {len(new_subtopics[:num_subtopics])}개 소제목 생성됨!")
                            st.rerun()
            
            with col_add:
                st.markdown("**직접 입력**")
                new_subtopic_name = st.text_input(
                    "소제목 이름",
                    placeholder="직접 입력하세요",
                    key="new_subtopic_empty"
                )
                if st.button("➕ 소제목 추가", key="add_subtopic_empty"):
                    if new_subtopic_name.strip():
                        chapter_data['subtopics'].append(new_subtopic_name)
                        chapter_data['subtopic_data'][new_subtopic_name] = {'questions': [], 'answers': [], 'content': ''}
                        st.success(f"'{new_subtopic_name}' 추가됨!")
                        st.rerun()

    # ====== 본문 작성 탭 하단: 작성된 본문 통합 보기 ======
    st.markdown("---")
    st.markdown("### 📖 작성된 본문 통합 보기")
    
    # 🔧 수정: 책 형식으로 깔끔하게 표시
    all_content_display = ""
    content_count_tab4 = 0
    
    # outline 순서대로 처리
    for ch_idx, ch in enumerate(st.session_state['outline'], 1):
        if ch in st.session_state['chapters']:
            ch_data = st.session_state['chapters'][ch]
            if 'subtopic_data' in ch_data:
                chapter_has_content = False
                chapter_content_parts = []
                
                # 소제목이 있는 경우
                subtopic_list = ch_data.get('subtopics', [])
                # 에필로그 등 소제목 없이 챕터 자체가 키인 경우도 포함
                if not subtopic_list and ch in ch_data['subtopic_data']:
                    subtopic_list = [ch]
                
                for st_name in subtopic_list:
                    st_data = ch_data['subtopic_data'].get(st_name, {})
                    if st_data.get('content'):
                        content_text = st_data['content']
                        chapter_content_parts.append(f"**{st_name}**\n\n{content_text}")
                        content_count_tab4 += 1
                        chapter_has_content = True
                
                if chapter_has_content:
                    # 챕터 제목은 한 번만, 소제목들은 그 아래에
                    all_content_display += f"\n\n---\n\n## {ch}\n\n"
                    all_content_display += "\n\n".join(chapter_content_parts)
    
    # 🔧 수정: 통일된 글자 수 계산 함수 사용
    pure_content = get_all_content_text()
    
    if pure_content:
        total_chars_tab4 = calculate_char_count(pure_content)
        st.success(f"✅ 총 {content_count_tab4}개 소제목 작성 완료 | {total_chars_tab4:,}자")
        
        with st.expander("📖 전체 본문 펼쳐보기", expanded=False):
            # 마크다운 형식으로 깔끔하게 표시
            for ch_idx, ch in enumerate(st.session_state['outline'], 1):
                if ch in st.session_state['chapters']:
                    ch_data = st.session_state['chapters'][ch]
                    if 'subtopic_data' in ch_data:
                        subtopic_list = ch_data.get('subtopics', [])
                        if not subtopic_list and ch in ch_data['subtopic_data']:
                            subtopic_list = [ch]
                        
                        chapter_has_content = False
                        chapter_contents = []
                        
                        for st_name in subtopic_list:
                            st_data = ch_data['subtopic_data'].get(st_name, {})
                            if st_data.get('content'):
                                # 본문 정제
                                cleaned_content = clean_content_for_display(st_data['content'], st_name, ch)
                                if cleaned_content.strip():
                                    chapter_contents.append((st_name, cleaned_content))
                                    chapter_has_content = True
                        
                        if chapter_has_content:
                            # 챕터 제목
                            st.markdown(f"## {ch}")
                            st.markdown("---")
                            
                            # 소제목과 본문
                            for st_name, content in chapter_contents:
                                st.markdown(f"**{st_name}**")
                                st.markdown(content)
                                st.markdown("")  # 빈 줄
    else:
        st.info("💡 아직 작성된 본문이 없습니다. 위에서 소제목을 선택하고 본문을 작성해주세요.")

# === TAB 5: 문체 다듬기 ===
with tabs[4]:
    st.markdown("## 문체 다듬기 & 품질 검사")
    
    # 작성된 본문이 있는지 확인
    has_content = False
    for ch_data in st.session_state['chapters'].values():
        if 'subtopic_data' in ch_data:
            for st_data in ch_data['subtopic_data'].values():
                if st_data.get('content'):
                    has_content = True
                    break
    
    if not has_content:
        st.info("💡 먼저 본문을 작성해주세요. 또는 아래에서 직접 텍스트를 입력할 수 있습니다.")
        
        direct_content = st.text_area(
            "다듬을 텍스트 직접 입력",
            height=300,
            placeholder="다듬고 싶은 텍스트를 여기에 붙여넣으세요..."
        )
        
        if direct_content:
            has_content = True
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown('<p class="section-label">Style</p>', unsafe_allow_html=True)
        st.markdown("### 문체 다듬기")
        
        # 챕터/소제목 선택 (콘텐츠가 있는 경우)
        content_options = []
        for ch in st.session_state['outline']:
            if ch in st.session_state['chapters']:
                ch_data = st.session_state['chapters'][ch]
                if 'subtopic_data' in ch_data:
                    for st_name, st_data in ch_data['subtopic_data'].items():
                        if st_data.get('content'):
                            content_options.append(f"{ch} > {st_name}")
        
        if content_options:
            selected_content = st.selectbox(
                "다듬을 콘텐츠 선택",
                content_options,
                key="refine_select"
            )
        
        style = st.selectbox(
            "목표 스타일",
            ["친근한", "전문적", "직설적", "스토리텔링"],
            key="style_select"
        )
        
        if st.button("✨ 문체 다듬기", key="refine_btn"):
            content_to_refine = ""
            
            if content_options and selected_content:
                parts = selected_content.split(" > ")
                if len(parts) == 2:
                    ch, st_name = parts
                    content_to_refine = st.session_state['chapters'][ch]['subtopic_data'][st_name]['content']
            elif 'direct_content' in dir() and direct_content:
                content_to_refine = direct_content
            
            if content_to_refine:
                with st.spinner("다듬는 중..."):
                    refined = refine_content(content_to_refine, style)
                    st.session_state['refined_content'] = refined
            else:
                st.error("다듬을 콘텐츠를 선택해주세요.")
        
        if st.session_state.get('refined_content'):
            st.text_area("다듬어진 본문", value=st.session_state['refined_content'], height=400)
            
            if st.button("원본에 적용", key="apply_refined"):
                if content_options and selected_content:
                    parts = selected_content.split(" > ")
                    if len(parts) == 2:
                        ch, st_name = parts
                        st.session_state['chapters'][ch]['subtopic_data'][st_name]['content'] = st.session_state['refined_content']
                        trigger_auto_save()
                        st.success("적용됨!")
                        st.rerun()
    
    with col2:
        st.markdown('<p class="section-label">Quality</p>', unsafe_allow_html=True)
        st.markdown("### 품질 검사")
        
        if st.button("🔍 베스트셀러 체크", key="quality_btn"):
            content_to_check = ""
            
            if content_options and selected_content:
                parts = selected_content.split(" > ")
                if len(parts) == 2:
                    ch, st_name = parts
                    content_to_check = st.session_state['chapters'][ch]['subtopic_data'][st_name]['content']
            elif 'direct_content' in dir() and direct_content:
                content_to_check = direct_content
            
            if content_to_check:
                with st.spinner("분석 중..."):
                    quality_result = check_quality(content_to_check)
                    st.session_state['quality_result'] = quality_result
            else:
                st.error("검사할 콘텐츠를 선택해주세요.")
        
        if st.session_state.get('quality_result'):
            st.markdown(f"""
            <div class="info-card">
                {st.session_state['quality_result'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)

# === TAB 6: 최종 출력 ===
with tabs[5]:
    st.markdown("## 최종 출력 & 마케팅")
    
    col1, col2 = st.columns([1.5, 1])
    
    with col1:
        st.markdown('<p class="section-label">Export</p>', unsafe_allow_html=True)
        st.markdown("### 전자책 다운로드")
        
        # 제목/부제 입력
        book_title = st.text_input("전자책 제목", value=st.session_state.get('book_title', ''), key="final_title")
        subtitle = st.text_input("부제", value=st.session_state.get('subtitle', ''), key="final_subtitle")
        
        st.session_state['book_title'] = book_title
        st.session_state['subtitle'] = subtitle
        
        # 스타일 설정
        st.markdown("### 스타일 설정")
        
        col_s1, col_s2 = st.columns(2)
        with col_s1:
            font_family = st.selectbox("본문 폰트", ["Pretendard", "Noto Sans KR", "Nanum Gothic"], key="font_select")
            font_size = st.selectbox("본문 크기", ["16px", "17px", "18px"], key="fontsize_select")
        with col_s2:
            line_height = st.selectbox("줄간격", ["1.8", "1.9", "2.0"], key="lineheight_select")
            max_width = st.selectbox("최대 폭", ["700px", "800px", "900px"], key="maxwidth_select")
        
        # 세부 설정
        with st.expander("상세 설정"):
            title_size = st.selectbox("제목 크기", ["32px", "36px", "40px"], key="titlesize_select")
            chapter_size = st.selectbox("챕터 제목 크기", ["24px", "26px", "28px"], key="chaptersize_select")
            subtopic_size = st.selectbox("소제목 크기", ["18px", "20px", "22px"], key="subtopicsize_select")
            text_color = st.color_picker("본문 색상", "#333333", key="textcolor_select")
        
        st.markdown("---")
        
        # 전체 책 내용 생성
        full_book_txt = ""
        full_book_html = ""
        
        if book_title:
            full_book_txt += f"{book_title}\n"
            full_book_html += f"<h1>{book_title}</h1>\n"
        if subtitle:
            full_book_txt += f"{subtitle}\n"
            full_book_html += f"<p style='color: #666; font-size: 14px; margin-top: -10px;'>{subtitle}</p>\n"
        
        full_book_txt += "\n" + "="*50 + "\n\n"
        full_book_html += "<hr style='border: none; border-top: 1px solid #ddd; margin: 30px 0;'>\n"
        
        # 챕터별 내용 수집
        for chapter in st.session_state['outline']:
            if chapter in st.session_state['chapters']:
                ch_data = st.session_state['chapters'][chapter]
                
                if 'subtopic_data' in ch_data:
                    chapter_has_content = False
                    for st_name in ch_data.get('subtopics', []):
                        st_data = ch_data['subtopic_data'].get(st_name, {})
                        if st_data.get('content'):
                            chapter_has_content = True
                            break
                    
                    if chapter_has_content:
                        full_book_txt += f"\n{chapter}\n" + "-"*40 + "\n\n"
                        full_book_html += f"<h2 style='font-size: {chapter_size}; margin-top: 50px;'>{chapter}</h2>\n"
                        
                        for st_name in ch_data.get('subtopics', []):
                            st_data = ch_data['subtopic_data'].get(st_name, {})
                            if st_data.get('content'):
                                full_book_txt += f"\n{st_name}\n\n{st_data['content']}\n\n"
                                
                                paragraphs = st_data['content'].split('\n\n')
                                full_book_html += f"<h3 style='font-size: {subtopic_size}; margin-top: 35px;'>{st_name}</h3>\n"
                                for para in paragraphs:
                                    para = para.strip()
                                    if para:
                                        full_book_html += f"<p style='font-size: {font_size}; line-height: {line_height}; color: {text_color};'>{para}</p>\n"
        
        # HTML 문서 완성
        html_content = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{book_title or '전자책'}</title>
    <link href="https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css" rel="stylesheet">
    <style>
        @page {{
            margin: 2cm;
        }}
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        h1 {{
            font-size: {title_size};
            color: #111;
            margin-bottom: 10px;
            font-weight: 700;
        }}
        h2 {{
            font-size: {chapter_size};
            color: #222;
            margin-top: 50px;
            margin-bottom: 20px;
            font-weight: 700;
        }}
        h3 {{
            font-size: {subtopic_size};
            color: #333;
            margin-top: 35px;
            margin-bottom: 15px;
            font-weight: 700;
        }}
        body {{
            font-family: '{font_family}', sans-serif;
            max-width: {max_width};
            margin: 0 auto;
            padding: 60px 20px;
            word-break: keep-all;
            font-weight: 500;
        }}
    </style>
</head>
<body>
{full_book_html}
</body>
</html>"""
        
        # 워드 파일 생성 함수
        def create_docx():
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # 제목
                if book_title:
                    title_para = doc.add_paragraph()
                    title_run = title_para.add_run(book_title)
                    title_run.font.size = Pt(28)
                    title_run.font.bold = True
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 부제
                if subtitle:
                    sub_para = doc.add_paragraph()
                    sub_run = sub_para.add_run(subtitle)
                    sub_run.font.size = Pt(14)
                    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if book_title or subtitle:
                    doc.add_paragraph()  # 빈 줄
                
                # 본문
                for chapter in st.session_state['outline']:
                    if chapter in st.session_state['chapters']:
                        ch_data = st.session_state['chapters'][chapter]
                        
                        if 'subtopic_data' in ch_data:
                            chapter_has_content = False
                            for st_name in ch_data.get('subtopics', []):
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    chapter_has_content = True
                                    break
                            
                            if chapter_has_content:
                                # 챕터 제목
                                ch_para = doc.add_paragraph()
                                ch_run = ch_para.add_run(chapter)
                                ch_run.font.size = Pt(20)
                                ch_run.font.bold = True
                                
                                for st_name in ch_data.get('subtopics', []):
                                    st_data = ch_data['subtopic_data'].get(st_name, {})
                                    if st_data.get('content'):
                                        # 소제목
                                        st_para = doc.add_paragraph()
                                        st_run = st_para.add_run(st_name)
                                        st_run.font.size = Pt(14)
                                        st_run.font.bold = True
                                        
                                        # 본문
                                        paragraphs = st_data['content'].split('\n\n')
                                        for para in paragraphs:
                                            para = para.strip()
                                            if para:
                                                p = doc.add_paragraph()
                                                run = p.add_run(para)
                                                run.font.size = Pt(11)
                
                # 파일로 저장
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer.getvalue()
            except ImportError:
                return None
        
        # 다운로드 버튼들
        col_dl1, col_dl2 = st.columns(2)
        with col_dl1:
            st.download_button(
                "📄 TXT 다운로드",
                full_book_txt,
                file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col_dl2:
            st.download_button(
                "🌐 HTML 다운로드",
                html_content,
                file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.html",
                mime="text/html",
                use_container_width=True
            )
        
        col_dl3, col_dl4 = st.columns(2)
        with col_dl3:
            # 워드 파일 다운로드
            docx_data = create_docx()
            if docx_data:
                st.download_button(
                    "📘 워드(DOCX) 다운로드",
                    docx_data,
                    file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            else:
                st.info("워드 파일: python-docx 필요")
        
        with col_dl4:
            # 한글 파일은 RTF로 대체 (hwp는 라이브러리 제한)
            # RTF 유니코드 지원 헤더
            rtf_content = "{\\rtf1\\ansi\\ansicpg949\\deff0\n"
            rtf_content += "{\\fonttbl{\\f0\\fnil\\fcharset129 맑은 고딕;}}\n"
            rtf_content += "\\f0\\fs24\n"
            rtf_content += escape_rtf_unicode(book_title or '') + "\\par\n"
            rtf_content += escape_rtf_unicode(subtitle or '') + "\\par\n"
            rtf_content += "\\par\n"
            
            for chapter in st.session_state['outline']:
                if chapter in st.session_state['chapters']:
                    ch_data = st.session_state['chapters'][chapter]
                    if 'subtopic_data' in ch_data:
                        chapter_has_content = any(ch_data['subtopic_data'].get(st_name, {}).get('content') for st_name in ch_data.get('subtopics', []))
                        if chapter_has_content:
                            rtf_content += "\\par\\b " + escape_rtf_unicode(chapter) + "\\b0\\par\\par\n"
                            for st_name in ch_data.get('subtopics', []):
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    rtf_content += "\\b " + escape_rtf_unicode(st_name) + "\\b0\\par\n"
                                    rtf_content += escape_rtf_unicode(st_data['content']) + "\\par\\par\n"
            rtf_content += "}"
            
            st.download_button(
                "📗 RTF 다운로드 (한글호환)",
                rtf_content.encode('utf-8'),
                file_name=f"{book_title or 'ebook'}_{datetime.now().strftime('%Y%m%d')}.rtf",
                mime="application/rtf",
                use_container_width=True
            )
        
        st.caption("💡 RTF 파일은 한글, 워드, 리브레오피스 등에서 열 수 있습니다.")
        
        # 미리보기 버튼
        st.markdown("---")
        
        # 📖 작성된 전체 본문 종합 보기
        st.markdown("### 📖 작성된 본문 종합 보기")
        
        all_content = ""
        content_count = 0
        
        # outline 순서대로 처리 (목차 순서 보장) - 책 형식
        for ch_idx, chapter in enumerate(st.session_state['outline'], 1):
            if chapter in st.session_state['chapters']:
                ch_data = st.session_state['chapters'][chapter]
                if 'subtopic_data' in ch_data:
                    chapter_has_content = False
                    chapter_content_parts = []
                    
                    # 소제목 목록 가져오기
                    subtopic_list = ch_data.get('subtopics', [])
                    # 에필로그 등 소제목 없이 챕터 자체가 키인 경우
                    if not subtopic_list and chapter in ch_data['subtopic_data']:
                        subtopic_list = [chapter]
                    
                    for st_name in subtopic_list:
                        st_data = ch_data['subtopic_data'].get(st_name, {})
                        if st_data.get('content'):
                            chapter_content_parts.append(f"**{st_name}**\n\n{st_data['content']}")
                            content_count += 1
                            chapter_has_content = True
                    
                    if chapter_has_content:
                        # 챕터 제목은 한 번만, 소제목들은 그 아래에
                        all_content += f"\n\n---\n\n## {chapter}\n\n"
                        all_content += "\n\n".join(chapter_content_parts)
        
        if all_content:
            st.success(f"✅ 총 {content_count}개 소제목 작성 완료")
            
            # 🔧 수정: 통일된 글자 수 계산
            pure_content_tab6 = get_all_content_text()
            total_chars = calculate_char_count(pure_content_tab6)
            st.caption(f"📊 총 {total_chars:,}자 / 약 {total_chars//500}페이지 (500자/페이지 기준)")
            
            # 본문 표시 - 마크다운 형식
            with st.expander("📖 전체 본문 펼쳐보기", expanded=False):
                for ch_idx, chapter in enumerate(st.session_state['outline'], 1):
                    if chapter in st.session_state['chapters']:
                        ch_data = st.session_state['chapters'][chapter]
                        if 'subtopic_data' in ch_data:
                            subtopic_list = ch_data.get('subtopics', [])
                            if not subtopic_list and chapter in ch_data['subtopic_data']:
                                subtopic_list = [chapter]
                            
                            chapter_has_content = False
                            chapter_contents = []
                            
                            for st_name in subtopic_list:
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    # 본문 정제
                                    cleaned_content = clean_content_for_display(st_data['content'], st_name, chapter)
                                    if cleaned_content.strip():
                                        chapter_contents.append((st_name, cleaned_content))
                                        chapter_has_content = True
                            
                            if chapter_has_content:
                                # 챕터 제목
                                st.markdown(f"## {chapter}")
                                st.markdown("---")
                                
                                # 소제목과 본문
                                for st_name, content in chapter_contents:
                                    st.markdown(f"**{st_name}**")
                                    st.markdown(content)
                                    st.markdown("")  # 빈 줄
            
            # 편집 가능한 텍스트 영역
            with st.expander("✏️ 전체 본문 편집하기 (텍스트)", expanded=False):
                # 편집용 텍스트 생성
                edit_text = ""
                for chapter in st.session_state['outline']:
                    if chapter in st.session_state['chapters']:
                        ch_data = st.session_state['chapters'][chapter]
                        if 'subtopic_data' in ch_data:
                            subtopic_list = ch_data.get('subtopics', [])
                            if not subtopic_list and chapter in ch_data['subtopic_data']:
                                subtopic_list = [chapter]
                            
                            chapter_has_content = False
                            for st_name in subtopic_list:
                                st_data = ch_data['subtopic_data'].get(st_name, {})
                                if st_data.get('content'):
                                    if not chapter_has_content:
                                        edit_text += f"\n\n{'='*50}\n{chapter}\n{'='*50}\n\n"
                                        chapter_has_content = True
                                    # 본문 정제
                                    cleaned = clean_content_for_display(st_data['content'], st_name, chapter)
                                    edit_text += f"[{st_name}]\n\n{cleaned}\n\n"
                
                edited_all = st.text_area(
                    "전체 본문 (편집 가능)",
                    value=edit_text.strip(),
                    height=600,
                    key="full_content_edit"
                )
                st.caption("여기서 수정한 내용은 개별 소제목에는 반영되지 않습니다. 최종 다운로드용으로만 사용됩니다.")
        else:
            st.info("💡 아직 작성된 본문이 없습니다. '④ 본문 작성' 탭에서 먼저 본문을 작성해주세요.")
        
        st.markdown("---")
        
        if st.button("👁️ 스타일 미리보기", key="preview_btn", use_container_width=True):
            st.session_state['show_preview'] = True
        
        # HTML 미리보기
        if st.session_state.get('show_preview'):
            st.markdown("### 스타일 미리보기")
            preview_sample = f"""
            <div style="font-family: '{font_family}', sans-serif; max-width: {max_width}; line-height: {line_height}; color: {text_color}; font-size: {font_size}; border: 1px solid #ddd; padding: 30px; border-radius: 8px; background: #fff;">
                <h1 style="font-size: {title_size}; font-weight: 700; color: #111; margin-bottom: 5px;">{book_title or '전자책 제목'}</h1>
                <p style="color: #666; font-size: 14px;">{subtitle or '부제목'}</p>
                <hr style="border: none; border-top: 1px solid #ddd; margin: 20px 0;">
                <h2 style="font-size: {chapter_size}; font-weight: 700; color: #222;">챕터1: 왜 열심히 하는 사람이 가난할까</h2>
                <h3 style="font-size: {subtopic_size}; font-weight: 700; color: #333;">그날 통장 잔고 47만원</h3>
                <p style="font-size: {font_size}; line-height: {line_height};">2019년 3월. 통장 잔고를 확인했습니다. 47만원. 월급날까지 2주. 저는 바닥이었습니다.</p>
                <p style="font-size: {font_size}; line-height: {line_height};">솔직히 말씀드리면, 저도 처음엔 몰랐습니다. 열심히만 하면 되는 줄 알았거든요. 새벽 6시에 일어나서 밤 11시까지 일했습니다. 주말도 없었습니다.</p>
            </div>
            """
            st.markdown(preview_sample, unsafe_allow_html=True)
    
    with col2:
        st.markdown('<p class="section-label">Marketing</p>', unsafe_allow_html=True)
        st.markdown("### 마케팅 카피")
        
        if st.button("카피 생성하기", key="marketing_btn"):
            with st.spinner("생성 중..."):
                marketing = generate_marketing_copy(
                    st.session_state.get('book_title', st.session_state['topic']),
                    st.session_state.get('subtitle', ''),
                    st.session_state['topic'],
                    st.session_state['target_persona']
                )
                st.session_state['marketing_copy'] = marketing
        
        if st.session_state.get('marketing_copy'):
            st.markdown(f"""
            <div class="info-card">
                {st.session_state['marketing_copy'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)

# --- 자동 저장 처리 ---
if st.session_state.get('auto_save_trigger'):
    st.session_state['auto_save_trigger'] = False
    auto_save_data = get_auto_save_data()
    auto_save_json = json.dumps(auto_save_data, ensure_ascii=False, indent=2)
    file_name = st.session_state.get('book_title', '전자책') or '전자책'
    file_name = re.sub(r'[^\w\s가-힣-]', '', file_name)[:20]
    
    st.toast("💾 자동 저장 준비됨!")
    
    # 사이드바에 자동 저장 다운로드 표시
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 🔔 자동 저장")
        st.download_button(
            "💾 백업 다운로드",
            auto_save_json,
            file_name=f"자동저장_{file_name}_{datetime.now().strftime('%H%M')}.json",
            mime="application/json",
            use_container_width=True,
            type="primary"
        )
        st.caption("중요 작업 완료됨 - 백업 권장!")

# --- 푸터 ---
st.markdown("""
<div class="premium-footer">
    <span class="premium-footer-text">전자책 작성 프로그램 — </span><span class="premium-footer-author">남현우 작가</span>
</div>
""", unsafe_allow_html=True)
